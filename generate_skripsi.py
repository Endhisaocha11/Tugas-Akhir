# -*- coding: utf-8 -*-
"""
Generator Skripsi PawfectCare — Smart Cat Feeder IoT  (Versi Lengkap v2)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── PAGE SETUP ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height   = Cm(29.7)
section.page_width    = Cm(21.0)
section.left_margin   = Cm(4.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ── HELPERS ──────────────────────────────────────────────────────────────────
def set_run_font(run, size=12, bold=False, italic=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_heading(text, size=14, bold=True, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p

def add_body(text, indent=True, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p

def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=11, italic=True)

def add_page_break():
    doc.add_page_break()

def add_table_styled(headers, rows, caption=''):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=11, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=11)
    if caption:
        add_caption(caption)
    return table

def add_diagram(lines, caption=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run('\n'.join(lines))
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    if caption:
        add_caption(caption)

def add_code(lines, caption=''):
    """Blok kode firmware — Courier New 9pt, rata kiri"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run('\n'.join(lines))
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    if caption:
        add_caption(caption)

# ═══════════════════════════════════════════════════════════════════════════
#  HALAMAN JUDUL
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run(
    'PERANCANGAN DAN IMPLEMENTASI SISTEM PEMBERI PAKAN KUCING\n'
    'OTOMATIS BERBASIS INTERNET OF THINGS (IoT)\n'
    'DENGAN PEMANTAUAN REAL-TIME MENGGUNAKAN ESP32 DAN FIREBASE'
)
set_run_font(r, size=16, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('(PawfectCare Smart Cat Feeder)')
set_run_font(r2, size=14, bold=True, italic=True)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('SKRIPSI')
set_run_font(r3, size=14, bold=True)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Diajukan untuk memenuhi sebagian persyaratan\nmemperoleh gelar Sarjana Teknik Informatika')
set_run_font(r4, size=12)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('Oleh:\n[NAMA MAHASISWA]\nNIM: [XXXXX]')
set_run_font(r5, size=12, bold=True)

for _ in range(3):
    doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run(
    'PROGRAM STUDI TEKNIK INFORMATIKA\n'
    'FAKULTAS [TEKNIK / ILMU KOMPUTER]\n'
    'UNIVERSITAS [...]\n'
    '[KOTA], 2025'
)
set_run_font(r6, size=12, bold=True)
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  ABSTRAK
# ═══════════════════════════════════════════════════════════════════════════
add_heading('ABSTRAK', size=14, center=True)
add_body(
    'PawfectCare adalah sistem pemberi pakan kucing otomatis berbasis IoT yang '
    'mengintegrasikan dua perangkat ESP32 (Feeder 1 dan Feeder 2), platform Firebase, '
    'dan aplikasi web React+TypeScript. Feeder 1 berfungsi sebagai unit dispensasi utama '
    'dengan sensor berat HX711, servo motor, sensor ultrasonik HC-SR04, dan OLED SH1106. '
    'Feeder 2 berfungsi sebagai unit monitoring tambahan dengan sensor HC-SR04 dan OLED SSD1306.'
)
add_body(
    'Data sensor dikirim secara real-time ke Firebase Realtime Database dan Firestore. '
    'Sistem mendukung pemberian pakan manual maupun terjadwal (hingga 10 slot/hari), '
    'dengan kalkulasi kebutuhan kalori otomatis berbasis metode RER sesuai standar WSAVA.'
)
add_body(
    'Hasil pengujian: akurasi dispensasi rata-rata 1,0 gram dari target (toleransi ±1,5 g), '
    'latensi data ESP32-ke-dashboard rata-rata di bawah 2 detik, dan jadwal otomatis '
    'berjalan dengan toleransi 0-60 detik dari waktu yang ditetapkan.'
)
doc.add_paragraph()
p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = p_kw.add_run('Kata Kunci: ')
set_run_font(r1, size=12, bold=True)
r2 = p_kw.add_run('IoT, ESP32, Firebase, Smart Cat Feeder, HX711, Servo Motor, React, TypeScript')
set_run_font(r2, size=12, italic=True)
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  BAB I — PENDAHULUAN
# ═══════════════════════════════════════════════════════════════════════════
add_heading('BAB I', size=14, center=True, space_before=0)
add_heading('PENDAHULUAN', size=14, center=True, space_before=0)

add_heading('1.1 Latar Belakang', size=12, bold=True)
add_body(
    'Kucing merupakan hewan peliharaan yang membutuhkan pola makan teratur dan porsi yang '
    'terukur. Ketidakteraturan pemberian pakan dapat menyebabkan obesitas (memicu FLUTD, '
    'diabetes, arthritis) atau malnutrisi. Pemilik yang sibuk sering kesulitan menjaga '
    'konsistensi jadwal dan porsi makan kucing.'
)
add_body(
    'Perkembangan teknologi Internet of Things (IoT) memungkinkan pembangunan sistem '
    'pemberi pakan otomatis yang dapat dipantau dan dikontrol dari jarak jauh. Dengan '
    'memanfaatkan ESP32, Firebase, dan aplikasi web modern, sistem PawfectCare '
    'dirancang untuk menjawab permasalahan tersebut secara terpadu.'
)

add_heading('1.2 Rumusan Masalah', size=12, bold=True)
add_body('Berdasarkan latar belakang, rumusan masalah dalam penelitian ini adalah:')
add_number('Bagaimana merancang mekanisme dispensasi pakan yang akurat (toleransi ±1,5 gram) menggunakan servo motor dan sensor berat HX711 pada ESP32?')
add_number('Bagaimana membangun komunikasi real-time antara dua perangkat ESP32 dan aplikasi web melalui Firebase?')
add_number('Bagaimana menghitung kebutuhan kalori dan porsi pakan kucing secara otomatis berdasarkan profil biologis?')
add_number('Bagaimana merancang antarmuka web yang informatif untuk monitoring, kontrol, dan analitik pemberian pakan?')

add_heading('1.3 Tujuan Penelitian', size=12, bold=True)
add_number('Merancang dan mengimplementasikan sistem smart feeder berbasis ESP32 dengan akurasi dispensasi tinggi.')
add_number('Membangun komunikasi IoT real-time antara dua perangkat ESP32 dan Firebase.')
add_number('Mengembangkan algoritma kalkulasi nutrisi kucing berbasis standar veteriner (RER/WSAVA).')
add_number('Membuat aplikasi web monitoring dengan dashboard, analitik, riwayat, dan kontrol jadwal.')

add_heading('1.4 Batasan Masalah', size=12, bold=True)
add_bullet('Sistem hanya mendukung pakan kering (dry food/kibble).')
add_bullet('Satu perangkat Feeder 1 (ESP32_FEEDER_01) per akun admin.')
add_bullet('Jadwal pemberian pakan maksimal 10 slot per hari.')
add_bullet('Konektivitas menggunakan Wi-Fi 2.4 GHz.')
add_bullet('Perhitungan kalori menggunakan metode RER standar veteriner.')
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  BAB II — TINJAUAN PUSTAKA
# ═══════════════════════════════════════════════════════════════════════════
add_heading('BAB II', size=14, center=True, space_before=0)
add_heading('TINJAUAN PUSTAKA', size=14, center=True, space_before=0)

add_heading('2.1 Internet of Things (IoT)', size=12, bold=True)
add_body(
    'IoT adalah paradigma teknologi yang menghubungkan perangkat fisik ke internet untuk '
    'pertukaran data otomatis. Arsitektur IoT terdiri dari tiga lapisan: '
    '(1) lapisan persepsi (sensor/aktuator), (2) lapisan jaringan (protokol komunikasi), '
    'dan (3) lapisan aplikasi (cloud & antarmuka). Dalam PawfectCare: '
    'sensor ESP32 = lapisan persepsi; Wi-Fi + Firebase = lapisan jaringan; '
    'aplikasi web React = lapisan aplikasi.'
)

add_heading('2.2 ESP32 Microcontroller', size=12, bold=True)
add_body(
    'ESP32 (Espressif Systems) adalah SoC dual-core Xtensa LX6 240 MHz dengan '
    'Wi-Fi 802.11 b/g/n dan Bluetooth 4.2 terintegrasi. Cocok untuk proyek IoT '
    'karena tidak membutuhkan modul Wi-Fi tambahan dan memiliki banyak pin GPIO.'
)
add_table_styled(
    ['Parameter', 'Spesifikasi'],
    [
        ['Prosesor', 'Dual-core Xtensa LX6, 240 MHz'],
        ['Flash', '4 MB'],
        ['SRAM', '520 KB'],
        ['Wi-Fi', '802.11 b/g/n, 2.4 GHz'],
        ['GPIO', '34 pin'],
        ['I2C', '2 channel'],
        ['ADC', '12-bit, 18 channel'],
        ['Tegangan', '3.3V (beberapa pin toleran 5V)'],
        ['Daya (Wi-Fi aktif)', '~240 mA'],
    ],
    caption='Tabel 2.1 Spesifikasi ESP32'
)

add_heading('2.3 Sensor Ultrasonik HC-SR04', size=12, bold=True)
add_body(
    'HC-SR04 mengukur jarak menggunakan prinsip Time of Flight: memancarkan gelombang '
    '40 kHz via pin TRIG, lalu mengukur durasi pantulan di pin ECHO. '
    'Rumus: Jarak (cm) = (Durasi ECHO (us) x 0.034) / 2. '
    'Digunakan di kedua feeder untuk mengukur level stok pakan di dalam hopper.'
)
add_table_styled(
    ['Parameter', 'Feeder 1 (ESP32_FEEDER_01)', 'Feeder 2 (ESP32_FEEDER_02)'],
    [
        ['Pin TRIG', 'GPIO 12', 'GPIO 17'],
        ['Pin ECHO', 'GPIO 14', 'GPIO 5'],
        ['Jarak "PENUH"', '3.0 cm', '2.0 cm'],
        ['Jarak "KOSONG"', '17.0 cm', '13.0 cm'],
        ['Filter EMA', '0.7 x lama + 0.3 x baru', '0.7 x lama + 0.3 x baru'],
    ],
    caption='Tabel 2.2 Perbandingan Konfigurasi HC-SR04 Feeder 1 vs Feeder 2'
)

add_heading('2.4 Sensor Berat HX711 & Load Cell', size=12, bold=True)
add_body(
    'HX711 adalah ADC 24-bit yang dirancang untuk antarmuka load cell. '
    'Load cell menghasilkan sinyal tegangan mV saat mengalami beban; HX711 '
    'memperkuat dan mengonversinya ke nilai digital via pin DT (data) dan SCK (clock). '
    'Hanya digunakan di Feeder 1 (ESP32_FEEDER_01). '
    'Faktor kalibrasi aktual di firmware: 404.0 (dapat diupdate via Firebase RTDB).'
)
add_table_styled(
    ['Parameter', 'Spesifikasi'],
    [
        ['Resolusi ADC', '24-bit'],
        ['Pin DT', 'GPIO 32'],
        ['Pin SCK', 'GPIO 33'],
        ['Load Cell kapasitas', '5 kg'],
        ['Faktor kalibrasi', '404.0 (default firmware)'],
        ['Sample tare', '20 sample (scale.tare(20))'],
        ['Akurasi setelah kalibrasi', '+/- 0.5 gram'],
    ],
    caption='Tabel 2.3 Spesifikasi HX711 & Load Cell'
)

add_heading('2.5 Servo Motor', size=12, bold=True)
add_body(
    'Servo motor adalah aktuator rotasi dengan umpan balik posisi via sinyal PWM. '
    'Dalam sistem ini, servo membuka/menutup pintu hopper untuk mengalirkan pakan. '
    'Sudut 0 derajat = tertutup; sudut 20 derajat = terbuka maksimal (MAX_SERVO_ANGLE). '
    'Pemilihan 20 derajat (bukan 90 derajat) untuk mengontrol laju aliran pakan agar '
    'tidak terlalu deras dan memudahkan fase fine-tuning.'
)

add_heading('2.6 Display OLED', size=12, bold=True)
add_body(
    'Dua jenis OLED digunakan dalam proyek ini. Feeder 1 menggunakan SH1106 (1.3 inci, '
    '128x64 piksel, I2C, library Adafruit_SH110X) dengan antarmuka grafis kustom '
    'termasuk animasi kucing kawaii. Feeder 2 menggunakan SSD1306 (0.96 inci, '
    '128x64 piksel, I2C, library Adafruit_SSD1306) dengan tampilan teks sederhana.'
)
add_table_styled(
    ['Parameter', 'Feeder 1 — SH1106', 'Feeder 2 — SSD1306'],
    [
        ['Ukuran', '1.3 inci', '0.96 inci'],
        ['Resolusi', '128 x 64 piksel', '128 x 64 piksel'],
        ['Library', 'Adafruit_SH110X', 'Adafruit_SSD1306'],
        ['Alamat I2C', '0x3C', '0x3C'],
        ['Inisialisasi', 'display.begin(0x3C, true)', 'display.begin(SSD1306_SWITCHCAPVCC, 0x3C)'],
        ['Layar', '3 layar kustom (Main/Feeding/Done)', '1 layar teks status'],
    ],
    caption='Tabel 2.4 Perbandingan OLED Feeder 1 dan Feeder 2'
)

add_heading('2.7 Firebase Platform', size=12, bold=True)
add_body(
    'Firebase (Google) adalah Backend-as-a-Service (BaaS) yang digunakan sebagai '
    'backbone cloud sistem PawfectCare. Empat layanan dimanfaatkan secara bersamaan '
    'sesuai kebutuhan data masing-masing:'
)
add_table_styled(
    ['Layanan Firebase', 'Fungsi dalam Sistem', 'Siapa yang Mengakses'],
    [
        ['Authentication', 'Login email/password, dua role: SUPER_ADMIN dan USER', 'Web App + ESP32 (email khusus)'],
        ['Realtime Database (RTDB)', 'Data sensor real-time, perintah servo, jadwal, status online', 'ESP32 (tulis) + Web App (baca)'],
        ['Cloud Firestore', 'Profil kucing, riwayat pemberian pakan, profil user, history profil', 'Web App (baca+tulis)'],
        ['Storage', 'Foto profil kucing', 'Web App (upload/download)'],
    ],
    caption='Tabel 2.5 Layanan Firebase yang Digunakan'
)

add_heading('2.8 React & TypeScript', size=12, bold=True)
add_body(
    'Aplikasi web dibangun dengan React 18 + TypeScript menggunakan bundler Vite. '
    'Styling menggunakan Tailwind CSS, animasi dengan Motion/React, '
    'dan visualisasi data dengan library Recharts. '
    'TypeScript memberikan keamanan tipe statis yang penting untuk data IoT yang kompleks.'
)

add_heading('2.9 Perhitungan Kebutuhan Pakan (RER/WSAVA)', size=12, bold=True)
add_body(
    'Perhitungan kebutuhan kalori menggunakan metode RER (Resting Energy Requirement) '
    'yang direkomendasikan WSAVA (World Small Animal Veterinary Association):'
)
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_f.add_run('RER (kcal/hari) = 70 x (Berat Badan kg)^0.75')
r_f.font.name = 'Courier New'; r_f.font.size = Pt(12); r_f.font.bold = True

add_body('Kebutuhan energi harian (DER) = RER x Faktor Metabolisme (k):')
add_table_styled(
    ['Kondisi Kucing', 'Faktor k'],
    [
        ['Anak kucing (< 1 tahun)', '2.5'],
        ['Tidak steril, dewasa', '1.6'],
        ['Steril, dewasa, aktif', '1.4'],
        ['Steril, dewasa, tidak aktif', '1.2'],
        ['Kucing tua (> 7 tahun)', '1.1'],
        ['Kucing obesitas (BCS 4-5), target turun berat', '0.8'],
    ],
    caption='Tabel 2.6 Faktor Metabolisme (k)'
)
add_body('Target gram/hari = DER / kandungan kalori per gram pakan (kcal/g).')

add_heading('2.10 Penelitian Terkait', size=12, bold=True)
add_table_styled(
    ['Peneliti / Tahun', 'Teknologi', 'Kekurangan vs PawfectCare'],
    [
        ['Zhao et al. (2019)', 'Arduino + ESP8266 + Blynk', 'Tanpa timbangan presisi, tanpa analitik'],
        ['Maulana et al. (2021)', 'NodeMCU + MQTT + Android', 'UI terbatas Android, tanpa kalkulasi kalori'],
        ['Prasetyo et al. (2022)', 'Raspberry Pi + Firebase', 'Biaya perangkat mahal, tanpa dual-device'],
        ['Sistem ini (2025)', 'ESP32 x2 + Firebase + React', 'Dual-device, kalori otomatis, analitik lengkap'],
    ],
    caption='Tabel 2.7 Perbandingan dengan Penelitian Terkait'
)
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  BAB III — METODE PENELITIAN
# ═══════════════════════════════════════════════════════════════════════════
add_heading('BAB III', size=14, center=True, space_before=0)
add_heading('METODE PENELITIAN', size=14, center=True, space_before=0)

add_heading('3.1 Kerangka Penelitian', size=12, bold=True)
add_body(
    'Penelitian menggunakan metode Research and Development (R&D) dengan pendekatan '
    'prototype. Tahapan: studi literatur, analisis kebutuhan, perancangan, '
    'implementasi, pengujian, dan evaluasi.'
)
add_diagram([
    '+----------------------------------------------------------+',
    '|                  KERANGKA PENELITIAN                     |',
    '+----------------------------------------------------------+',
    '| [1] Studi Literatur & Analisis Kebutuhan                 |',
    '|           v                                              |',
    '| [2] Perancangan Sistem                                   |',
    '|     +-- Hardware (ESP32, sensor, servo)                  |',
    '|     +-- Firmware (Arduino/C++)                           |',
    '|     +-- Database (Firebase RTDB + Firestore)             |',
    '|     +-- Aplikasi Web (React + TypeScript)                |',
    '|           v                                              |',
    '| [3] Implementasi                                         |',
    '|           v                                              |',
    '| [4] Pengujian (Unit + Integrasi + Sistem)                |',
    '|           v                                              |',
    '| [5] Evaluasi & Dokumentasi                               |',
    '+----------------------------------------------------------+',
], caption='Gambar 3.1 Kerangka Penelitian')

add_heading('3.2 Arsitektur Sistem Keseluruhan', size=12, bold=True)
add_body(
    'Sistem PawfectCare terdiri dari dua perangkat ESP32 yang masing-masing '
    'berkomunikasi dengan Firebase, serta aplikasi web yang membaca data dari Firebase '
    'secara real-time. Feeder 1 adalah unit utama (dispensasi + monitoring); '
    'Feeder 2 adalah unit monitoring tambahan (stok pakan saja).'
)
add_diagram([
    '+=====================================================================+',
    '|                 ARSITEKTUR SISTEM PAWFECTCARE                       |',
    '+====================+==================+============================+',
    '| LAYER FISIK        | LAYER CLOUD      | LAYER APLIKASI             |',
    '+====================+==================+============================+',
    '|                    |                  |                            |',
    '| ESP32_FEEDER_01    |  Firebase RTDB   |  Aplikasi Web (React)      |',
    '| + HC-SR04  ------->|  /devices/       |  + Dashboard               |',
    '| + HX711           |  ESP32_FEEDER_01 |  + Feeding Control (Admin) |',
    '| + Servo    <-------|  ESP32_FEEDER_02 |  + Analytics               |',
    '| + OLED SH1106      |  (command, data) |  + Feeding History         |',
    '|                    |                  |  + Cat Profile             |',
    '| ESP32_FEEDER_02    |  Firestore       |  + Notifications           |',
    '| + HC-SR04  ------->|  /cats           |  + Education               |',
    '| + OLED SSD1306     |  /feedingLogs    |  + Device Settings         |',
    '|                    |  /users          |                            |',
    '|                    |  /devices        |  Firebase Auth             |',
    '|                    |  /catProfileHist |  + SUPER_ADMIN             |',
    '|                    |                  |  + USER (monitoring)       |',
    '+====================+==================+============================+',
], caption='Gambar 3.2 Arsitektur Sistem PawfectCare')

add_heading('3.3 Perbedaan Feeder 1 dan Feeder 2', size=12, bold=True)
add_body('Dua perangkat ESP32 memiliki peran berbeda dalam ekosistem PawfectCare:')
add_table_styled(
    ['Aspek', 'ESP32_FEEDER_01 (Feeder 1)', 'ESP32_FEEDER_02 (Feeder 2)'],
    [
        ['Peran', 'Dispensasi + Monitoring utama', 'Monitoring stok tambahan'],
        ['Sensor stok', 'HC-SR04 (TRIG=12, ECHO=14)', 'HC-SR04 (TRIG=17, ECHO=5)'],
        ['Sensor berat', 'HX711 + Load Cell (DT=32, SCK=33)', 'Tidak ada'],
        ['Aktuator', 'Servo Motor (GPIO 27)', 'Tidak ada'],
        ['Display', 'OLED SH1106 1.3" (3 layar animasi)', 'OLED SSD1306 0.96" (teks status)'],
        ['Library OLED', 'Adafruit_SH110X', 'Adafruit_SSD1306'],
        ['Terima perintah', 'Ya (command node di RTDB)', 'Tidak'],
        ['Jadwal otomatis', 'Ya (10 slot/hari)', 'Tidak'],
        ['onDisconnect', 'setBoolOnDisconnect (RTDB SDK)', 'setOffline() manual'],
        ['Rentang stok', 'fullDist=3cm, emptyDist=17cm', 'fullDist=2cm, emptyDist=13cm'],
    ],
    caption='Tabel 3.1 Perbandingan Feeder 1 dan Feeder 2'
)

add_heading('3.4 Perancangan Perangkat Keras', size=12, bold=True)
add_heading('3.4.1 Konfigurasi Pin Feeder 1 (ESP32_FEEDER_01)', size=11, bold=True)
add_table_styled(
    ['Komponen', 'Pin Komponen', 'Pin ESP32', 'Keterangan'],
    [
        ['HC-SR04', 'TRIG', 'GPIO 12', 'Output trigger ultrasonik'],
        ['HC-SR04', 'ECHO', 'GPIO 14', 'Input echo ultrasonik'],
        ['Servo Motor', 'Signal (PWM)', 'GPIO 27', 'Kontrol buka-tutup hopper'],
        ['HX711', 'DT', 'GPIO 32', 'Data serial load cell'],
        ['HX711', 'SCK', 'GPIO 33', 'Clock serial load cell'],
        ['OLED SH1106', 'SDA', 'GPIO 21', 'I2C data display'],
        ['OLED SH1106', 'SCL', 'GPIO 22', 'I2C clock display'],
    ],
    caption='Tabel 3.2 Pin ESP32_FEEDER_01'
)

add_heading('3.4.2 Konfigurasi Pin Feeder 2 (ESP32_FEEDER_02)', size=11, bold=True)
add_table_styled(
    ['Komponen', 'Pin Komponen', 'Pin ESP32', 'Keterangan'],
    [
        ['HC-SR04', 'TRIG', 'GPIO 17', 'Output trigger ultrasonik'],
        ['HC-SR04', 'ECHO', 'GPIO 5', 'Input echo ultrasonik'],
        ['OLED SSD1306', 'SDA', 'GPIO 21', 'I2C data display'],
        ['OLED SSD1306', 'SCL', 'GPIO 22', 'I2C clock display'],
    ],
    caption='Tabel 3.3 Pin ESP32_FEEDER_02'
)

add_diagram([
    '       SKEMA KONEKSI FEEDER 1 (ESP32_FEEDER_01)',
    '',
    '   HC-SR04        ESP32           Servo Motor',
    '   +-------+     +---------+      +----------+',
    '   |TRIG   |---->|GPIO 12  |      |Signal    |',
    '   |ECHO   |<----|GPIO 14  |----->|          |',
    '   |VCC 5V |--+  |GPIO 27  |      +----------+',
    '   |GND    |--+  |         |',
    '   +-------+  |  |GPIO 21  |----> OLED SH1106 SDA',
    '              |  |GPIO 22  |----> OLED SH1106 SCL',
    '   HX711      |  |GPIO 32  |----> HX711 DT',
    '   +-------+  |  |GPIO 33  |----> HX711 SCK',
    '   |DT     |--+  |5V / GND |--+',
    '   |SCK    |  +--| all GND |  |',
    '   |VCC    |--+  +---------+  |',
    '   |GND    |-------------------+',
    '   +-------+',
], caption='Gambar 3.3 Skema Koneksi Feeder 1')

add_heading('3.5 Perancangan Firmware Feeder 1', size=12, bold=True)
add_heading('3.5.1 Konstanta dan Parameter Firmware', size=11, bold=True)
add_body(
    'Berikut adalah konstanta-konstanta penting yang mendefinisikan perilaku '
    'sistem dispensasi di Feeder 1:'
)
add_table_styled(
    ['Konstanta', 'Nilai Aktual', 'Penjelasan'],
    [
        ['MAX_SERVO_ANGLE', '20 derajat', 'Sudut buka maksimal servo (kontrol laju aliran)'],
        ['PAN_WEIGHT_G', '3.0 g', 'Estimasi berat wadah pakan untuk field weightGross'],
        ['DISPENSE_BUFFER_G', '7.0 g', 'Servo ditutup 7g sebelum target (kompensasi pakan melayang)'],
        ['DISPENSE_TOLERANCE_G', '1.5 g', 'Toleransi akhir yang dianggap berhasil'],
        ['DISPENSE_PULSE_MS', '80 ms', 'Durasi buka servo per-pulse fase fine tuning'],
        ['DISPENSE_SETTLE_MS', '2500 ms', 'Waktu tunggu pakan settling ke timbangan'],
        ['DISPENSE_MAX_PULSES', '6 kali', 'Maksimal pulse fase fine sebelum menyerah'],
        ['loadCellFactor', '404.0 (default)', 'Faktor kalibrasi HX711, dapat diperbarui dari Firebase'],
    ],
    caption='Tabel 3.4 Konstanta Firmware Feeder 1 (nilai aktual dari kode)'
)

add_heading('3.5.2 Algoritma Pembacaan Stok Pakan (readFoodStock)', size=11, bold=True)
add_body(
    'Fungsi readFoodStock() membaca jarak dari sensor HC-SR04, '
    'mengonversinya ke persentase stok, lalu menerapkan filter EMA '
    '(Exponential Moving Average) untuk mengurangi noise:'
)
add_code([
    'int readFoodStock() {',
    '  // Trigger HC-SR04',
    '  digitalWrite(TRIG_PIN, LOW);  delayMicroseconds(2);',
    '  digitalWrite(TRIG_PIN, HIGH); delayMicroseconds(10);',
    '  digitalWrite(TRIG_PIN, LOW);',
    '',
    '  long dur = pulseIn(ECHO_PIN, HIGH, 30000);',
    '  if (dur <= 0) return last;       // timeout: kembalikan nilai terakhir',
    '',
    '  float dist = dur * 0.034f / 2.0f;  // konversi durasi -> jarak (cm)',
    '  foodDistance = dist;',
    '',
    '  // Konversi jarak -> persentase (fullDist=3cm, emptyDist=17cm)',
    '  const float fullDist = 3.0f, emptyDist = 17.0f;',
    '  dist    = constrain(dist, fullDist, emptyDist);',
    '  float pct = ((emptyDist - dist) / (emptyDist - fullDist)) * 100.0f;',
    '  pct = constrain(pct, 0.0f, 100.0f);',
    '',
    '  // EMA: bobot 70% nilai lama + 30% nilai baru',
    '  pct  = last * 0.7f + pct * 0.3f;',
    '  last = (int)pct;',
    '  return last;',
    '}',
], caption='Kode 3.1 Fungsi readFoodStock() Feeder 1')

add_heading('3.5.3 Algoritma Pembacaan Berat (readWeight)', size=11, bold=True)
add_body(
    'Fungsi readWeight() membaca berat dari HX711 dengan 1 sample, '
    'menghilangkan nilai negatif dan noise kecil (< 1 gram):'
)
add_code([
    'float readWeight() {',
    '  if (!scale.is_ready()) return foodWeight;  // HX711 belum siap',
    '',
    '  float raw = scale.get_units(1);   // 1 sample cepat',
    '',
    '  if (raw < 0.0f) raw = 0.0f;       // hilangkan negatif (drift)',
    '  if (raw < 1.0f) raw = 0.0f;       // threshold noise < 1g dianggap 0',
    '',
    '  return roundf(raw * 10.0f) / 10.0f;  // bulatkan ke 1 desimal',
    '}',
    '',
    '// readWeightFast() -- digunakan selama feeding loop (1 sample, tanpa rounding)',
    'float readWeightFast() {',
    '  return scale.get_units(1);',
    '}',
], caption='Kode 3.2 Fungsi readWeight() dan readWeightFast() Feeder 1')

add_heading('3.5.4 Algoritma Dispensasi Pakan Dua Fase (executeFeed)', size=11, bold=True)
add_body(
    'Algoritma dispensasi menggunakan dua fase untuk mencapai akurasi tinggi. '
    'Fase 1 (Bulk): servo dibuka penuh untuk mengalirkan pakan dengan cepat '
    'hingga mendekati target. Fase 2 (Fine Tuning): servo dibuka singkat berulang '
    'untuk mencapai target dengan presisi tinggi.'
)
add_diagram([
    '   FLOWCHART executeFeed(targetGrams)',
    '',
    '   START: terima target (gram)',
    '   Firebase: servoStatus = "active", command/status = "executing"',
    '          v',
    '   init0 = baca berat awal (10 sample)',
    '          v',
    '   +----------------------------------------------+',
    '   |  FASE 1: BULK DISPENSING                     |',
    '   |  Loop:                                       |',
    '   |    servo.write(20)  -- buka 80ms             |',
    '   |    servo.write(0)   -- tutup                 |',
    '   |    delay(1200ms)    -- tunggu settling        |',
    '   |    dispensed = baca_berat - init0             |',
    '   |    update animasi OLED                       |',
    '   |    heartbeat Firebase (isOnline=true)         |',
    '   |    BREAK jika dispensed >= target - 7g        |',
    '   |    BREAK jika dispensed > target x 2 (safety) |',
    '   +----------------------------------------------+',
    '          v',
    '   Settling: baca ulang dengan 5 sample',
    '          v',
    '   +----------------------------------------------+',
    '   |  FASE 2: FINE TUNING (max 6 pulse)           |',
    '   |  Ulangi hingga 6x atau target tercapai:      |',
    '   |    if dispensed >= target - 1.5g: BREAK       |',
    '   |    servo.write(20) -- buka 80ms              |',
    '   |    servo.write(0)  -- tutup                  |',
    '   |    delay(2500ms)   -- tunggu settling panjang |',
    '   |    dispensed = baca_berat - init0             |',
    '   +----------------------------------------------+',
    '          v',
    '   targetMet = (dispensed >= target - 1.5g)',
    '          v',
    '   Firebase: servoStatus = "idle"',
    '   Firebase: command/status = "done" atau "incomplete"',
    '   Firebase: weight, weightGross diperbarui',
    '          v',
    '   jika targetMet: tampilkan animasi "Nyam!" OLED (10x, 300ms)',
    '          v',
    '   END',
], caption='Gambar 3.4 Flowchart Algoritma Dispensasi Dua Fase')

add_code([
    'void executeFeed(float targetGrams) {',
    '  String base = "/devices/" + String(DEVICE_ID);',
    '  Firebase.RTDB.setString(&fbdo, (base+"/servoStatus").c_str(), "active");',
    '  Firebase.RTDB.setString(&fbdo, (base+"/command/status").c_str(), "executing");',
    '',
    '  float init0    = scale.get_units(10);  // berat awal (10 sample)',
    '  float dispensed = 0;',
    '  bool  targetMet = false;',
    '',
    '  // ---- FASE 1: BULK ----------------------------------------',
    '  while (dispensed < targetGrams - DISPENSE_BUFFER_G) {',
    '    feederServo.write(MAX_SERVO_ANGLE);  // buka servo',
    '    delay(80);                           // DISPENSE_PULSE_MS',
    '    feederServo.write(0);                // tutup servo',
    '    delay(1200);                         // settling bulk',
    '    dispensed = readWeightFast() - init0;',
    '    if (dispensed < 0) dispensed = 0;',
    '    animFrame++;',
    '    drawFeedingScreen(dispensed, targetGrams);',
    '    // Heartbeat agar web tidak anggap offline',
    '    Firebase.RTDB.setBool(&fbdo, (base+"/isOnline").c_str(), true);',
    '    if (dispensed > targetGrams * 2) break;  // safety',
    '  }',
    '',
    '  // Baca ulang setelah settling',
    '  dispensed = scale.get_units(5) - init0;',
    '  if (dispensed < 0) dispensed = 0;',
    '',
    '  // ---- FASE 2: FINE TUNING ---------------------------------',
    '  for (int pulse = 0; pulse < DISPENSE_MAX_PULSES; pulse++) {',
    '    if (dispensed >= targetGrams - DISPENSE_TOLERANCE_G) {',
    '      targetMet = true; break;',
    '    }',
    '    feederServo.write(MAX_SERVO_ANGLE);',
    '    delay(DISPENSE_PULSE_MS);   // 80ms',
    '    feederServo.write(0);',
    '    delay(DISPENSE_SETTLE_MS);  // 2500ms',
    '    dispensed = readWeightFast() - init0;',
    '    if (dispensed < 0) dispensed = 0;',
    '    animFrame++;',
    '    drawFeedingScreen(dispensed, targetGrams);',
    '  }',
    '  if (!targetMet && dispensed >= targetGrams - DISPENSE_TOLERANCE_G)',
    '    targetMet = true;',
    '',
    '  Firebase.RTDB.setString(&fbdo, (base+"/servoStatus").c_str(), "idle");',
    '  Firebase.RTDB.setString(&fbdo, (base+"/command/status").c_str(),',
    '                           targetMet ? "done" : "incomplete");',
    '  foodWeight = readWeight();',
    '  Firebase.RTDB.setFloat(&fbdo, (base+"/weight").c_str(), foodWeight);',
    '  Firebase.RTDB.setFloat(&fbdo, (base+"/weightGross").c_str(), foodWeight + PAN_WEIGHT_G);',
    '  if (targetMet) {',
    '    for (int i = 0; i < 10; i++) { drawFeedDoneScreen(dispensed); delay(300); }',
    '  }',
    '}',
], caption='Kode 3.3 Implementasi Lengkap executeFeed()')

add_heading('3.5.5 Mekanisme Jadwal Otomatis (checkScheduledFeeding)', size=11, bold=True)
add_body(
    'Fungsi checkScheduledFeeding() dipanggil setiap 1 detik dari loop utama. '
    'Fungsi ini membaca node schedule dari Firebase RTDB, mencocokkan waktu sekarang '
    'dengan slot yang aktif, dan mengeksekusi pemberian pakan jika cocok. '
    'Array execSlots[] mencatat slot yang sudah dieksekusi hari ini untuk mencegah '
    'eksekusi ganda dalam satu hari.'
)
add_code([
    'void checkScheduledFeeding() {',
    '  // Baca jadwal dari Firebase RTDB',
    '  String path = "/devices/" + String(DEVICE_ID) + "/schedule";',
    '  if (!Firebase.RTDB.getJSON(&fbdo, path.c_str())) return;',
    '',
    '  // Cek enabled dan dailyLimitReached',
    '  // Iterasi setiap slot dalam array slots[]',
    '  // Jika slot.active && slot.time == waktuSekarang && belum dieksekusi hari ini:',
    '    //   executeFeed(slot.amount)',
    '    //   tambah slot ke execSlots[]',
    '    //   tulis scheduledFeedLog ke RTDB (akan dibaca web app)',
    '',
    '  // scheduledFeedLog yang ditulis ke RTDB:',
    '  // { amount, slot (HH:MM), ts (unix ms), processed: false }',
    '  // Web app mendeteksi processed=false -> buat Firestore log -> set processed=true',
    '}',
], caption='Kode 3.4 Alur checkScheduledFeeding() (disederhanakan)')

add_heading('3.5.6 Deteksi Status Online (onDisconnect)', size=11, bold=True)
add_body(
    'Feeder 1 menggunakan fitur setBoolOnDisconnect dari Firebase RTDB SDK '
    'untuk secara otomatis mengatur isOnline=false saat koneksi ESP32 terputus. '
    'Ini lebih andal dibanding polling karena dieksekusi langsung di server Firebase.'
)
add_code([
    'void registerOnDisconnect() {',
    '  String path = "/devices/" + String(DEVICE_ID) + "/isOnline";',
    '  // Server Firebase akan set isOnline=false jika koneksi ESP32 terputus',
    '  if (Firebase.RTDB.setBoolOnDisconnect(&fbdo, path.c_str(), false)) {',
    '    Serial.println("onDisconnect registered OK");',
    '  }',
    '}',
    '',
    '// Dipanggil setelah Firebase.begin() dan setelah setiap reconnect:',
    '// initFirebase() -> registerOnDisconnect()',
    '// loop() -> if Firebase reconnect -> registerOnDisconnect()',
], caption='Kode 3.5 Registrasi onDisconnect Feeder 1')

add_heading('3.5.7 Tampilan OLED Feeder 1 (SH1106 128x64)', size=11, bold=True)
add_body(
    'Feeder 1 memiliki tiga layar OLED yang masing-masing dirancang piksel-per-piksel '
    'menggunakan library Adafruit_GFX. Berikut layout piksel layar utama (idle):'
)
add_diagram([
    '  LAYOUT OLED SH1106 128x64px (Main Screen)',
    '',
    '  x=0                    x=60           x=127',
    '  +----------------------+---------------+  y=0',
    '  | WiFi | PAWFECT       | HH:MM         |  y=1 (header)',
    '  +----------------------------------------------+  y=10 (separator)',
    '  |                      |  FOOD STOCK           |  y=11',
    '  |   KUCING KAWAII      |  [can][=====] XX%     |  y=20',
    '  |   (lingkaran r=17)   |  ---- divider ----    |  y=32',
    '  |   Telinga y~10       |  FOOD WEIGHT          |  y=34',
    '  |   Kumis, hidung      |  [bowl] XX g          |  y=44',
    '  |   Paws y=57-63       |                       |  y=55',
    '  +----------------------+-----------------------+  y=63',
    '',
    '  Layar Feeding Screen:',
    '  Kiri : kucing happy (mata ^ ^)',
    '  Kanan: animasi dispenser + pellet jatuh + bar stok + berat real-time',
    '',
    '  Layar Feed Done Screen:',
    '  Kucing besar tengah + mangkuk + teks "Nyam~ XXg!" + sparkle',
], caption='Gambar 3.5 Layout Piksel OLED SH1106')

add_heading('3.6 Perancangan Firmware Feeder 2', size=12, bold=True)
add_body(
    'Feeder 2 (ESP32_FEEDER_02) memiliki firmware yang lebih sederhana karena '
    'hanya berfungsi sebagai unit monitoring stok pakan tambahan. '
    'Tidak ada servo, HX711, atau logika dispensasi.'
)
add_diagram([
    '  FLOWCHART FIRMWARE FEEDER 2 (setup)',
    '',
    '  START -> initOLED() -> connectWiFi() -> initNTP() -> initFirebase()',
    '  -> Tampil "SYSTEM READY" -> masuk loop()',
    '',
    '  LOOP (setiap iterasi):',
    '  +------------------------------------------+',
    '  | WiFi putus? -> setOffline() -> reconnect  |',
    '  |                                           |',
    '  | Tiap 1000ms: readFoodStock()               |',
    '  |   HC-SR04 TRIG=17, ECHO=5                 |',
    '  |   fullDist=2cm, emptyDist=13cm            |',
    '  |   EMA: 0.7*lama + 0.3*baru               |',
    '  |                                           |',
    '  | Tiap 2000ms: sendDataToFirebase()          |',
    '  |   isOnline, foodStock, distance_cm, time  |',
    '  |                                           |',
    '  | Tiap 500ms: update tampilan OLED SSD1306  |',
    '  |   Baris 1: "PawfectCare"                  |',
    '  |   Baris 2: DEVICE_ID                     |',
    '  |   Baris 3: waktu NTP (HH:MM:SS)          |',
    '  |   Baris 4: Stock: XX%                    |',
    '  |   Baris 5: Dist: X.X cm                  |',
    '  +------------------------------------------+',
], caption='Gambar 3.6 Flowchart Firmware Feeder 2')

add_body(
    'Perbedaan penanganan status online: Feeder 2 tidak dapat menggunakan '
    'setBoolOnDisconnect karena library Firebase_ESP_Client berbasis HTTP (bukan '
    'WebSocket). Sebagai gantinya, fungsi setOffline() dipanggil secara eksplisit '
    'sebelum reconnect:'
)
add_code([
    '// Feeder 2 -- tidak ada setBoolOnDisconnect (HTTP-based library)',
    'void setOffline() {',
    '  if (!Firebase.ready()) return;',
    '  String path = "/devices/" + String(DEVICE_ID) + "/isOnline";',
    '  Firebase.RTDB.setBool(&fbdo, path.c_str(), false);',
    '}',
    '',
    '// Dipanggil di loop() sebelum reconnect:',
    '// if (WiFi.status() != WL_CONNECTED) { setOffline(); connectWiFi(); return; }',
    '// if (!Firebase.ready()) { setOffline(); Firebase.begin(...); ... }',
], caption='Kode 3.6 Penanganan Status Online Feeder 2')

add_heading('3.7 Perancangan Firebase Realtime Database (RTDB)', size=12, bold=True)
add_body(
    'Firebase RTDB digunakan untuk data yang membutuhkan latensi sangat rendah '
    '(< 100ms): telemetri sensor real-time, perintah kontrol servo, jadwal, '
    'dan status online. Data disusun di bawah node /devices/{deviceId}/.'
)
add_diagram([
    'Firebase Realtime Database — Struktur Node',
    '',
    '/devices/',
    '  +-- ESP32_FEEDER_01/',
    '  |     +-- isOnline: true/false',
    '  |     +-- foodStock: 75          <- % stok (HC-SR04)',
    '  |     +-- weight: 42.5           <- berat mangkuk (HX711, gram)',
    '  |     +-- weightGross: 45.5      <- berat + wadah pan',
    '  |     +-- distance_cm: 6.3       <- jarak ultrasonik mentah',
    '  |     +-- time: "14:30:00"       <- waktu NTP dari ESP32',
    '  |     +-- servoStatus: "idle"    <- idle / active',
    '  |     +-- calibration/',
    '  |     |     +-- loadCellFactor: 404.0',
    '  |     +-- command/',
    '  |     |     +-- type: "feed"     <- feed / test_servo',
    '  |     |     +-- amount: 50       <- gram',
    '  |     |     +-- status: "pending"<- pending/executing/done/incomplete',
    '  |     +-- schedule/',
    '  |     |     +-- enabled: true',
    '  |     |     +-- dailyLimitReached: false',
    '  |     |     +-- slots: [{time:"07:00", amount:40, active:true}, ...]',
    '  |     +-- scheduledFeedLog/',
    '  |           +-- amount: 40',
    '  |           +-- slot: "07:00"',
    '  |           +-- ts: 1748000000000  <- unix ms',
    '  |           +-- processed: false   <- web app set true setelah buat log',
    '  |',
    '  +-- ESP32_FEEDER_02/',
    '        +-- isOnline: true/false',
    '        +-- foodStock: 60',
    '        +-- distance_cm: 7.1',
    '        +-- time: "14:30:00"',
], caption='Gambar 3.7 Struktur Firebase Realtime Database')

add_heading('3.8 Perancangan Cloud Firestore', size=12, bold=True)
add_body(
    'Firestore digunakan untuk data struktural yang membutuhkan query kompleks: '
    'profil kucing, riwayat pemberian pakan, profil pengguna, dan snapshot profil. '
    'Firestore mendukung query dengan filter (where), limit, dan realtime listener '
    'via onSnapshot yang digunakan di aplikasi web.'
)
add_diagram([
    'Cloud Firestore — Struktur Collections',
    '',
    '/users/{uid}',
    '  +-- uid, email',
    '  +-- role: "SUPER_ADMIN" | "USER"',
    '  +-- onboardingCompleted: boolean',
    '  +-- claimedDeviceId: string | null',
    '  +-- createdAt: timestamp',
    '',
    '/devices/{deviceId}',
    '  +-- isClaimed: boolean',
    '  +-- claimedBy: uid | null',
    '  +-- claimedAt: timestamp | null',
    '  +-- linkedCatId: string | null',
    '  +-- deviceName: string',
    '  +-- deviceType: string',
    '  (catatan: data telemetri live di RTDB, bukan di sini)',
    '',
    '/cats/{catId}',
    '  +-- ownerId: uid',
    '  +-- name, gender, age, weight, photoUrl',
    '  +-- isSterilized: boolean',
    '  +-- bodyCondition: 1-5',
    '  +-- kiloCaloriesPerKg: number',
    '  +-- dailyCalorieTarget: number',
    '  +-- dailyGramTarget: number',
    '  +-- feedingSchedule: [{time, amount, active}]',
    '  +-- profileUpdatedAt: timestamp',
    '  +-- dailyLimitReachedDate: "YYYY-MM-DD" | null',
    '  +-- dailyLimitResetDate: "YYYY-MM-DD" | null',
    '  +-- dailyAdjustments: {date, manualTotal, slots:[]}',
    '  +-- smartFeedEnabled: boolean',
    '  +-- activity: string',
    '',
    '/feedingLogs/{logId}',
    '  +-- catId, deviceId',
    '  +-- timestamp: unix ms',
    '  +-- amountRequested, amountDispensed: number (gram)',
    '  +-- weightBefore, weightAfter: number',
    '  +-- status: "success" | "failed" | "warning"',
    '  +-- notes: "manual" | "auto" | "scheduled"',
    '',
    '/catProfileHistory/{snapshotId}',
    '  +-- catId, ownerId',
    '  +-- savedAt, endedAt: timestamp',
    '  +-- [semua field CatProfile saat snapshot dibuat]',
], caption='Gambar 3.8 Struktur Cloud Firestore')

add_heading('3.9 Mekanisme Sinkronisasi Data di Aplikasi Web (useCatData)', size=12, bold=True)
add_body(
    'File useCatData.ts adalah custom hook React yang menjadi pusat data seluruh '
    'aplikasi. Hook ini menggabungkan data dari dua sumber berbeda: '
    'Firestore (metadata statis) dan RTDB (telemetri real-time). '
    'Tiga useEffect berjalan paralel untuk mengambil data dari sumber yang berbeda.'
)
add_diagram([
    '  ARSITEKTUR DATA FLOW — useCatData.ts',
    '',
    '  +-------------------+    onSnapshot    +---------------------+',
    '  | Firestore         |----------------->| Effect 1:           |',
    '  | /cats             |                 | cats, devices,      |',
    '  | /devices          |                 | profileHistory      |',
    '  | /catProfileHistory|                 +---------------------+',
    '  +-------------------+                         |',
    '                                                 v',
    '  +-------------------+    onValue       +---------------------+',
    '  | RTDB              |----------------->| Effect 2:           |',
    '  | /devices/{id}     |                 | rtdbDeviceData      |',
    '  | (telemetri live)  |                 | isOnline, weight,   |',
    '  | /devices/{id}/    |                 | servo, time         |',
    '  | scheduledFeedLog  |                 +---------------------+',
    '  +-------------------+                         |',
    '                                                 v',
    '  +-------------------+    onSnapshot    +---------------------+',
    '  | Firestore         |----------------->| Effect 3:           |',
    '  | /feedingLogs      |                 | feedingLogs[]       |',
    '  +-------------------+                 +---------------------+',
    '                                                 |',
    '                                                 v',
    '  +---------------------------------------------------+',
    '  | MERGE: Firestore metadata + RTDB telemetri        |',
    '  | isOnline = rtdb.isOnline && isRtdbLive (< 2 menit)|',
    '  | foodStockLevel = rtdb.foodStock                   |',
    '  | currentWeightOnScale = rtdb.weight                |',
    '  | servoStatus = rtdb.servoStatus                    |',
    '  +---------------------------------------------------+',
    '                    |',
    '                    v',
    '  Return: { cat, device, feedingLogs, profileHistory }',
    '  -> dikonsumsi Dashboard, Analytics, FeedingHistory, dll.',
], caption='Gambar 3.9 Alur Data useCatData.ts')

add_heading('3.9.1 Mekanisme Deteksi Status Online (isRtdbLive)', size=11, bold=True)
add_body(
    'Untuk mendeteksi apakah ESP32 benar-benar online, aplikasi web tidak hanya '
    'bergantung pada field isOnline di RTDB (karena nilainya bisa tidak berubah '
    'setelah ESP32 crash). Sebaliknya, ditambahkan pengecekan freshness data '
    'berdasarkan field time yang dikirim ESP32 via NTP:'
)
add_code([
    '// isRtdbLive: true hanya jika data RTDB dikirim dalam 2 menit terakhir',
    'const isRtdbLive = (() => {',
    '  void staleCheck;  // reactive trigger ticker 30 detik',
    '  const timeStr = rtdbDeviceData?.time;  // "HH:MM:SS" dari ESP32',
    '  if (!timeStr) return false;',
    '  try {',
    '    const [h, m, s] = timeStr.split(":").map(Number);',
    '    const now    = new Date();',
    '    const devSec = h * 3600 + m * 60 + s;',
    '    const nowSec = now.getHours()*3600 + now.getMinutes()*60 + now.getSeconds();',
    '    let diffSec  = Math.abs(nowSec - devSec);',
    '    if (diffSec > 43200) diffSec = 86400 - diffSec;  // tengah malam crossover',
    '    return diffSec < 120;  // toleransi 2 menit (120 detik)',
    '  } catch { return false; }',
    '})();',
    '',
    '// Device dianggap online hanya jika ketiganya terpenuhi:',
    '// 1. rtdbDeviceData.isOnline === true (ESP32 set aktif)',
    '// 2. isRtdbLive === true (data < 2 menit yang lalu)',
    'const isOnline = rtdbDeviceData?.isOnline === true && isRtdbLive;',
], caption='Kode 3.7 Mekanisme Deteksi isRtdbLive di useCatData.ts')

add_heading('3.9.2 Sinkronisasi Log Jadwal Otomatis', size=11, bold=True)
add_body(
    'Ketika ESP32 mengeksekusi jadwal otomatis, ia menulis scheduledFeedLog ke RTDB '
    'dengan processed=false. Web app mendeteksi ini via listener onValue, '
    'membuat entry Firestore feedingLog dengan notes="scheduled", '
    'lalu mengubah processed=true untuk mencegah duplikasi:'
)
add_code([
    '// ESP32 menulis ke RTDB setelah eksekusi jadwal:',
    '// /devices/ESP32_FEEDER_01/scheduledFeedLog/',
    '// { amount:40, slot:"07:00", ts:1748000000000, processed:false }',
    '',
    '// Web app (useCatData.ts) listener:',
    'const unsubSchedLog = onValue(schedLogRef, async (snap) => {',
    '  if (!snap.exists()) return;',
    '  const data = snap.val();',
    '  if (data.processed !== false || !data.amount) return;',
    '',
    '  // Set processed=true DULUAN -- cegah double-log',
    '  await set(ref(rtdb, `devices/${id}/scheduledFeedLog/processed`), true);',
    '',
    '  // Buat Firestore feedingLog dengan notes="scheduled"',
    '  await addDoc(collection(db, "feedingLogs"), {',
    '    catId:           catRef.current.id,',
    '    deviceId:        claimedDeviceId,',
    '    timestamp:       data.ts ?? Date.now(),',
    '    amountRequested: data.amount,',
    '    amountDispensed: data.amount,',
    '    status:          "success",',
    '    notes:           "scheduled",',
    '  });',
    '});',
], caption='Kode 3.8 Sinkronisasi Jadwal Otomatis RTDB -> Firestore')

add_heading('3.10 Perancangan Aplikasi Web — Alur Navigasi', size=12, bold=True)
add_diagram([
    '                ALUR NAVIGASI APLIKASI WEB',
    '',
    '  [Auth Screen (Login/Register)]',
    '         |',
    '         v login berhasil',
    '  [Cek Role Pengguna]',
    '         |                     |',
    '    SUPER_ADMIN               USER',
    '         |                     |',
    '  [Belum klaim device?]   [Monitoring Selection]',
    '         |                     |',
    '    Ya-> [Device Claim Screen] |',
    '         |                     |',
    '  [OnboardingFlow: isi profil kucing]',
    '         |',
    '         v selesai onboarding',
    '  +---------------------------------------------------------+',
    '  |                DASHBOARD LAYOUT                         |',
    '  +---------------------------------------------------------+',
    '  | Dashboard     | Analytics     | Feeding History         |',
    '  | Feeding Ctrl* | Cat Profile*  | Education               |',
    '  | Notifications | Device Set.*  | User Settings*          |',
    '  | (* hanya SUPER_ADMIN)                                   |',
    '  +---------------------------------------------------------+',
], caption='Gambar 3.10 Alur Navigasi Aplikasi Web')

add_heading('3.11 Perhitungan Kebutuhan Nutrisi', size=12, bold=True)
add_table_styled(
    ['Parameter', 'Nilai Contoh', 'Keterangan'],
    [
        ['Nama Kucing', 'Mochi', '—'],
        ['Berat Badan', '4.5 kg', 'Diinput pengguna'],
        ['Usia', '3 tahun', 'Dewasa'],
        ['Status Sterilisasi', 'Ya', '—'],
        ['Body Condition Score', '3 (Ideal)', 'Skala 1-5'],
        ['Tingkat Aktivitas', 'Rendah', 'Indoor, tidak aktif'],
        ['RER', '70 x 4.5^0.75 = 218 kcal', 'Resting Energy Requirement'],
        ['Faktor k', '1.2', 'Steril, dewasa, tidak aktif'],
        ['DER (Target Kalori)', '218 x 1.2 = 262 kcal/hari', 'Daily Energy Requirement'],
        ['Kandungan Pakan', '3500 kcal/kg = 3.5 kcal/g', 'Data label pakan (diinput user)'],
        ['Target Gram/Hari', '262 / 3.5 = 75 gram/hari', 'Hasil kalkulasi otomatis'],
    ],
    caption='Tabel 3.5 Contoh Perhitungan Kebutuhan Pakan Otomatis'
)

add_heading('3.12 Alat dan Bahan Penelitian', size=12, bold=True)
add_body(
    'Penelitian ini menggunakan kombinasi perangkat keras IoT, perangkat lunak firmware, '
    'layanan cloud, dan kerangka kerja aplikasi web. Rincian alat dan bahan beserta '
    'fungsinya dalam sistem dijabarkan pada subbab berikut.'
)

add_heading('3.12.1 Perangkat Keras (Alat)', size=11, bold=True)
add_table_styled(
    ['Komponen', 'Jumlah', 'Fungsi dalam Proyek'],
    [
        ['ESP32 Dev Board', '2 unit', 'Mikrokontroler utama tiap feeder; membaca sensor, mengirim telemetri ke Firebase RTDB, menerima perintah servo, dan menjalankan jadwal makan otomatis'],
        ['Sensor Ultrasonik HC-SR04', '2 unit (1/feeder)', 'Mengukur sisa stok pakan di hopper berbasis Time of Flight, dikonversi ke persentase stok'],
        ['Load Cell 5 kg + Modul HX711', '1 set (Feeder 1)', 'Menimbang mangkuk sebelum dan sesudah dispensasi untuk menghitung berat pakan aktual yang keluar (amountDispensed) dan estimasi pakan yang dimakan (amountEaten)'],
        ['Servo Motor SG90', '1 unit (Feeder 1)', 'Aktuator yang membuka/menutup pintu hopper (sudut 0-20 derajat) saat perintah feed diterima'],
        ['OLED SH1106 1.3"', '1 unit (Feeder 1)', 'Menampilkan animasi kucing dan status feeding langsung di perangkat fisik (3 layar)'],
        ['OLED SSD1306 0.96"', '1 unit (Feeder 2)', 'Menampilkan status teks sederhana (online/offline, persentase stok)'],
        ['Power Supply 5V', '2 unit', 'Catu daya untuk tiap ESP32 beserta sensor dan aktuator yang terhubung'],
        ['Housing/Hopper Pakan', '2 unit', 'Wadah fisik penyimpanan dan jalur dispensasi pakan kucing'],
    ],
    caption='Tabel 3.6 Alat (Perangkat Keras) yang Digunakan'
)

add_heading('3.12.2 Bahan — Perangkat Lunak Aplikasi Web', size=11, bold=True)
add_table_styled(
    ['Bahan / Library', 'Versi', 'Fungsi dalam Proyek'],
    [
        ['React + TypeScript', 'React 19.0.1, TypeScript 5.8.2', 'Membangun seluruh antarmuka web (Dashboard, FeedingControl, Analytics, dll) dengan keamanan tipe statis untuk struktur data IoT (CatProfile, DeviceStatus, FeedingLog)'],
        ['Vite', '6.2.3', 'Build tool dan dev server untuk kompilasi serta hot-reload aplikasi web selama pengembangan'],
        ['Tailwind CSS', '4.1.14', 'Utility-first CSS framework untuk styling seluruh komponen antarmuka'],
        ['Motion (motion/react)', '12.23.24', 'Animasi transisi modal, expand/collapse, dan hover effect pada antarmuka'],
        ['Recharts', '3.8.1', 'Visualisasi data berupa grafik (area chart, pie chart) pada Dashboard dan Analytics'],
        ['React Router DOM', '7.15.0', 'Navigasi antar halaman aplikasi (Dashboard, Control, History, Settings, dll)'],
        ['Lucide React', '0.546.0', 'Pustaka ikon yang digunakan di seluruh antarmuka aplikasi'],
        ['Firebase JS SDK', '12.13.0', 'Klien Firebase untuk Authentication, Firestore, Realtime Database, dan Storage dari sisi aplikasi web'],
    ],
    caption='Tabel 3.7 Bahan (Perangkat Lunak) Aplikasi Web yang Digunakan'
)

add_heading('3.12.3 Bahan — Perangkat Lunak Firmware ESP32', size=11, bold=True)
add_table_styled(
    ['Bahan / Library', 'Versi', 'Spesifikasi & Fungsi dalam Proyek'],
    [
        ['Arduino IDE / PlatformIO', 'IDE 2.3.x / PIO 6.x', 'Lingkungan pengembangan terintegrasi (IDE) tempat menulis, mengompilasi, dan mengunggah kode C++ firmware ke kedua ESP32. PlatformIO menyediakan manajemen dependensi otomatis melalui berkas platform.ini [61].'],
        ['Adafruit_SH110X', '2.1.x', 'Driver library untuk menggambar teks, grafik, dan animasi pada layar OLED SH1106/SH1107 milik Feeder 1. Bergantung pada library Adafruit_GFX sebagai lapisan antarmuka grafis. Digunakan untuk menampilkan tiga layar animasi kucing secara langsung pada perangkat fisik [60].'],
        ['Adafruit_SSD1306', '2.5.x', 'Driver library untuk layar OLED SSD1306 milik Feeder 2. Mendukung komunikasi I2C dan SPI. Digunakan untuk menampilkan teks status (online/offline, persentase stok) pada Feeder 2 [59].'],
        ['HX711 (bogde)', '0.7.5', 'Membaca sinyal milivolt mentah dari load cell via modul HX711 dan mengonversinya ke satuan gram menggunakan faktor kalibrasi (calibrationFactor) yang dapat diperbarui dari cloud melalui komponen DeviceSettings.tsx. Merupakan fondasi dari klaim akurasi +/-1,5 g pada proyek ini [36].'],
        ['ESP32Servo', '3.0.x', 'Mengirimkan sinyal PWM ke motor servo SG90 menggunakan timer perangkat keras ESP32 (bukan timer perangkat lunak) untuk mengatur sudut buka pintu hopper (0-20 derajat). Mendukung hingga 12 servo secara bersamaan [62].'],
        ['Firebase_ESP_Client (Mobizt)', '4.4.x', 'Jembatan komunikasi HTTPS dari ESP32 ke seluruh layanan Firebase (RTDB, Firestore, Auth, Storage). Digunakan untuk mengirimkan telemetri sensor (isOnline, foodStock, weight, servoStatus), menerima perintah makan/servo, serta menulis dan membaca calibrationFactor dari RTDB [43].'],
    ],
    caption='Tabel 3.8 Bahan (Perangkat Lunak) Firmware ESP32 yang Digunakan'
)

add_heading('3.12.4 Bahan — Layanan Cloud Firebase', size=11, bold=True)
add_table_styled(
    ['Layanan', 'Fungsi dalam Proyek'],
    [
        ['Authentication', 'Login email/password dan pembedaan role SUPER_ADMIN (dapat klaim device) vs USER (hanya monitoring)'],
        ['Realtime Database (RTDB)', 'Jalur data telemetri live dari ESP32: isOnline, foodStock, weight, servoStatus, calibrationFactor, di-listen real-time oleh aplikasi web'],
        ['Cloud Firestore', 'Penyimpanan data persisten/terstruktur: profil kucing, riwayat feeding, metadata klaim device, history perubahan profil'],
        ['Storage', 'Penyimpanan file foto profil kucing yang diunggah pengguna'],
    ],
    caption='Tabel 3.9 Bahan (Layanan Cloud) Firebase yang Digunakan'
)
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  BAB IV — IMPLEMENTASI DAN HASIL
# ═══════════════════════════════════════════════════════════════════════════
add_heading('BAB IV', size=14, center=True, space_before=0)
add_heading('IMPLEMENTASI DAN HASIL', size=14, center=True, space_before=0)

add_heading('4.1 Implementasi Perangkat Keras', size=12, bold=True)
add_table_styled(
    ['No.', 'Komponen', 'Jml', 'Digunakan di', 'Fungsi'],
    [
        ['1', 'ESP32 DevKitV1', '2', 'Feeder 1 & 2', 'Mikrokontroler utama + Wi-Fi'],
        ['2', 'HC-SR04', '2', 'Feeder 1 & 2', 'Monitor level stok pakan hopper'],
        ['3', 'HX711 + Load Cell 5kg', '1', 'Feeder 1', 'Timbangan berat pakan di mangkuk'],
        ['4', 'Servo Motor MG996R', '1', 'Feeder 1', 'Mekanisme buka-tutup dispensasi'],
        ['5', 'OLED SH1106 1.3"', '1', 'Feeder 1', 'Display animasi 3 layar'],
        ['6', 'OLED SSD1306 0.96"', '1', 'Feeder 2', 'Display teks status'],
        ['7', 'Power Supply 5V/2A', '2', 'Feeder 1 & 2', 'Catu daya sistem'],
        ['8', 'Wadah Pakan (Hopper)', '1', 'Feeder 1', 'Tempat penyimpanan kibble'],
        ['9', 'Mangkuk Stainless', '1', 'Feeder 1', 'Tempat makan (di atas load cell)'],
    ],
    caption='Tabel 4.1 Daftar Komponen Perangkat Keras'
)

add_heading('4.2 Implementasi Loop Utama Feeder 1', size=12, bold=True)
add_body(
    'Loop utama Feeder 1 menjalankan beberapa tugas periodik dengan timer independen '
    'menggunakan millis(). Semua tugas berjalan non-blocking agar tidak ada delay '
    'yang memblokir eksekusi lain:'
)
add_table_styled(
    ['Tugas', 'Interval', 'Fungsi yang Dipanggil'],
    [
        ['Cek perintah Firebase', '1.000 ms', 'checkCommand() -> executeFeed() atau executeTestServo()'],
        ['Cek jadwal otomatis', '1.000 ms', 'checkScheduledFeeding()'],
        ['Baca sensor', '250 ms', 'readFoodStock() + readWeight()'],
        ['Kirim data ke Firebase', '2.000 ms', 'sendDataToFirebase()'],
        ['Update tampilan OLED', '500 ms', 'drawMainScreen()'],
        ['Update waktu NTP', 'Setiap iterasi', 'timeClient.update()'],
    ],
    caption='Tabel 4.2 Tugas Periodik Loop Utama Feeder 1'
)

add_heading('4.3 Implementasi Data yang Dikirim ke Firebase RTDB', size=12, bold=True)
add_body('Setiap 2 detik, Feeder 1 mengirim paket data berikut ke Firebase RTDB:')
add_code([
    '// sendDataToFirebase() -- dikirim tiap 2 detik',
    'FirebaseJson json;',
    'json.set("isOnline",    true);',
    'json.set("foodStock",   foodStock);          // % (0-100)',
    'json.set("weight",      foodWeight);          // gram (setelah tare)',
    'json.set("weightGross", foodWeight + 3.0f);  // gram + berat wadah pan',
    'json.set("distance_cm", foodDistance);        // cm mentah HC-SR04',
    'json.set("time",        timeClient.getFormattedTime()); // "HH:MM:SS"',
    '',
    '// Feeder 2 mengirim subset yang lebih kecil:',
    'json.set("isOnline",    true);',
    'json.set("foodStock",   foodStock);',
    'json.set("distance_cm", foodDistance);',
    'json.set("time",        timeClient.getFormattedTime());',
], caption='Kode 4.1 Data yang Dikirim ke Firebase RTDB')

add_heading('4.4 Implementasi Kode Kunci Aplikasi Web', size=12, bold=True)
add_body(
    'Aplikasi web PawfectCare terdiri dari beberapa file TypeScript/React yang menjadi '
    'tulang punggung sistem. Berikut adalah file-file paling penting beserta '
    'penjelasan dan kode kuncinya:'
)

add_heading('4.4.1 firebase.ts — Inisialisasi Layanan Firebase', size=11, bold=True)
add_body(
    'File ini adalah titik masuk seluruh koneksi Firebase. Semua layanan '
    '(Auth, Firestore, RTDB, Storage) diinisialisasi di sini dan diekspor '
    'sebagai singleton yang digunakan oleh seluruh komponen aplikasi. '
    'Tanpa file ini tidak ada satu pun fitur yang dapat mengakses cloud.'
)
add_code([
    '// src/lib/firebase.ts',
    'import { getApps, initializeApp } from "firebase/app";',
    'import { getAuth }      from "firebase/auth";',
    'import { getFirestore } from "firebase/firestore";',
    'import { getDatabase }  from "firebase/database";',
    'import { getStorage }   from "firebase/storage";',
    'import firebaseConfig   from "../../firebase-applet-config.json";',
    '',
    'const app = initializeApp(firebaseConfig);',
    '',
    'export const auth    = getAuth(app);',
    'export const storage = getStorage(app);',
    '',
    '// Firestore -- mendukung custom database ID',
    'const db = (firebaseConfig as any).firestoreDatabaseId',
    '  ? getFirestore(app, (firebaseConfig as any).firestoreDatabaseId)',
    '  : getFirestore(app);',
    'export { db };',
    '',
    '// RTDB -- hanya diinisialisasi jika databaseURL tersedia',
    'let rtdb: any = null;',
    'if ((firebaseConfig as any).databaseURL) {',
    '  rtdb = getDatabase(app, (firebaseConfig as any).databaseURL);',
    '}',
    'export { rtdb };',
    '',
    '// secondaryApp -- untuk createUser tanpa logout admin saat ini',
    'export const secondaryApp  = getApps().find(a => a.name === "Secondary")',
    '  || initializeApp(firebaseConfig, "Secondary");',
    'export const secondaryAuth = getAuth(secondaryApp);',
], caption='Kode 4.2 firebase.ts — Inisialisasi Semua Layanan Firebase')

add_heading('4.4.2 AuthContext.tsx — Manajemen Autentikasi & Session', size=11, bold=True)
add_body(
    'AuthContext adalah React Context yang membungkus seluruh aplikasi. '
    'File ini menangani state login/logout, memuat profil pengguna dari Firestore, '
    'dan mengekspos data via hook useAuth(). '
    'Yang membuatnya penting secara teknis adalah penyelesaian 4 bug kritis:'
)
add_table_styled(
    ['Bug', 'Masalah', 'Solusi yang Diterapkan'],
    [
        ['#1 Memory Leak',
         'Listener Firestore (onSnapshot) di dalam callback onAuthStateChanged tidak pernah di-unsubscribe karena useEffect tidak bisa menangkap return value dari callback async.',
         'Simpan referensi unsubscribe ke useRef. Panggil cleanup eksplisit setiap kali auth state berubah.'],
        ['#2 Race Condition',
         'Jika user logout lalu login cepat, listener Firestore dari sesi lama masih aktif dan bisa overwrite data sesi baru.',
         'Cleanup listener lama sebelum mendaftarkan listener baru di setiap auth state change.'],
        ['#3 isAdmin Recalculation',
         'Nilai isAdmin dihitung ulang setiap render meski profile tidak berubah.',
         'Bungkus dengan useMemo agar hanya dihitung ulang saat profile berubah.'],
        ['#4 Context Re-render',
         'Object value context dibuat ulang setiap render, memaksa semua consumer re-render meski data sama.',
         'Bungkus context value dengan useMemo untuk mencegah re-render yang tidak perlu.'],
    ],
    caption='Tabel 4.7 Empat Bug Fix Kritis di AuthContext.tsx'
)
add_code([
    '// src/lib/AuthContext.tsx -- bagian kunci',
    '',
    '// FIX #1 & #2: simpan referensi listener Firestore ke useRef',
    'const unsubscribeProfileRef = useRef<(() => void) | null>(null);',
    '',
    'useEffect(() => {',
    '  const unsubscribeAuth = onAuthStateChanged(auth, (firebaseUser) => {',
    '    // Cleanup listener sesi sebelumnya (fix #2 race condition)',
    '    if (unsubscribeProfileRef.current) {',
    '      unsubscribeProfileRef.current();',
    '      unsubscribeProfileRef.current = null;',
    '    }',
    '    setUser(firebaseUser);',
    '    if (!firebaseUser) { setProfile(null); setLoading(false); return; }',
    '',
    '    // Daftarkan listener Firestore baru untuk profil user',
    '    const unsubProfile = onSnapshot(',
    '      doc(db, "users", firebaseUser.uid),',
    '      (snap) => {',
    '        setProfile(snap.exists() ? snap.data() as UserProfile : null);',
    '        setLoading(false);',
    '      },',
    '      () => { setProfile(null); setLoading(false); }',
    '    );',
    '    unsubscribeProfileRef.current = unsubProfile; // simpan referensi',
    '  });',
    '',
    '  return () => {',
    '    unsubscribeAuth();',
    '    unsubscribeProfileRef.current?.();  // cleanup saat unmount',
    '  };',
    '}, []);',
    '',
    '// FIX #3: isAdmin hanya dihitung ulang saat profile berubah',
    'const isAdmin = useMemo(',
    '  () => profile?.role === UserRole.SUPER_ADMIN,',
    '  [profile]',
    ');',
    '',
    '// FIX #4: context value tidak dibuat ulang setiap render',
    'const contextValue = useMemo<AuthContextType>(',
    '  () => ({ user, profile, loading, isAdmin }),',
    '  [user, profile, loading, isAdmin]',
    ');',
], caption='Kode 4.3 AuthContext.tsx — Solusi 4 Bug Kritis')

add_heading('4.4.3 App.tsx — Routing Berbasis Role', size=11, bold=True)
add_body(
    'App.tsx adalah komponen root yang mengimplementasikan state machine navigasi '
    'berbasis role pengguna. Setiap kondisi sistem (loading, tidak login, admin tanpa device, '
    'admin belum onboarding, user biasa) memiliki tampilan yang berbeda. '
    'Ini adalah implementasi Protected Route tanpa library routing eksternal.'
)
add_code([
    '// src/App.tsx -- logika routing utama (disederhanakan)',
    'function AppContent() {',
    '  const { user, profile, loading } = useAuth();',
    '',
    '  if (loading)  return <CatLoader />;           // Loading Firebase',
    '  if (!user)    return <AuthScreen />;           // Belum login',
    '  if (!profile) return <CatLoader />;           // Profil belum dimuat',
    '',
    '  const isAdmin = profile.role === UserRole.SUPER_ADMIN;',
    '',
    '  // ── ADMIN FLOW ─────────────────────────────────',
    '  // 1. Belum klaim device ESP32',
    '  if (isAdmin && !profile.claimedDeviceId)',
    '    return <DeviceSelectionScreen />;',
    '',
    '  // 2. Sudah klaim, belum setup profil kucing',
    '  if (isAdmin && !profile.onboardingCompleted)',
    '    return <OnboardingFlow isAdmin={true} />;',
    '',
    '  // 3. Siap -- tampilkan dashboard admin',
    '  if (isAdmin) return <DashboardLayout>...</DashboardLayout>;',
    '',
    '  // ── USER FLOW ──────────────────────────────────',
    '  const appMode = localStorage.getItem("appMode");',
    '  if (!appMode)           return <MonitoringSelection />;',
    '  if (appMode === "guest") return <OnboardingFlow isAdmin={false} />;',
    '  if (appMode === "monitor") return <DashboardLayout>...</DashboardLayout>;',
    '  return <MonitoringSelection />;',
    '}',
    '',
    '// Routing konten per tab (hanya di dalam DashboardLayout)',
    'function renderContent() {',
    '  switch (currentTab) {',
    '    case "dashboard":       return <Dashboard />;',
    '    case "feeding-control": return isAdmin ? <FeedingControl /> : <AccessDenied />;',
    '    case "cat-profile":     return <CatProfilePage />;',
    '    case "analytics":       return <Analytics />;',
    '    case "history":         return <FeedingHistory />;',
    '    case "education":       return <Education />;',
    '    case "notifications":   return <Notifications />;',
    '    case "settings":        return isAdmin ? <DeviceSettings /> : <AccessDenied />;',
    '    case "user-settings":   return isAdmin ? <UserSettings /> : <AccessDenied />;',
    '    default:                return <Dashboard />;',
    '  }',
    '}',
], caption='Kode 4.4 App.tsx — State Machine Routing Berbasis Role')

add_heading('4.4.4 useCatData.ts — Pusat Data Seluruh Aplikasi', size=11, bold=True)
add_body(
    'useCatData adalah custom React hook yang paling kritis dalam aplikasi. '
    'Hook ini menjalankan 3 listener paralel (Firestore cats + devices, '
    'RTDB telemetri, Firestore feedingLogs) lalu menggabungkan hasilnya menjadi '
    'satu objek data yang dikonsumsi oleh Dashboard, Analytics, History, '
    'FeedingControl, dan Notifications. '
    'Tanpa hook ini, tidak ada komponen yang bisa menampilkan data real-time.'
)
add_table_styled(
    ['Mekanisme', 'Penjelasan Teknis'],
    [
        ['3 useEffect paralel',
         'Effect 1: Firestore cats+devices+profileHistory. Effect 2: RTDB telemetri + listener scheduledFeedLog. Effect 3: Firestore feedingLogs dengan filter catId. Ketiganya berjalan independen dan update state masing-masing.'],
        ['isRtdbLive',
         'Membandingkan field "time" (HH:MM:SS) dari ESP32 dengan jam browser. Jika selisih > 120 detik, device dianggap offline meski isOnline=true. Ticker 30 detik (staleCheck) memaksa re-evaluasi berkala.'],
        ['Data fusion (merge)',
         'Device status akhir = gabungan Firestore claim metadata (isClaimed, deviceName) + RTDB telemetri (weight, stock, servo). isOnline hanya true jika rtdb.isOnline===true DAN isRtdbLive===true.'],
        ['scheduledFeedLog sync',
         'Listener RTDB di Effect 2 mendeteksi processed=false, langsung set true (cegah duplikasi), lalu buat Firestore feedingLog dengan notes="scheduled". ESP32 tidak perlu tahu Firestore.'],
        ['catRef (stale closure fix)',
         'useCatData menggunakan useRef untuk menyimpan cat terkini agar listener RTDB tidak punya stale closure saat mengakses catId.'],
    ],
    caption='Tabel 4.8 Mekanisme Kunci useCatData.ts'
)
add_code([
    '// src/lib/useCatData.ts -- bagian paling kritis',
    '',
    '// Effect 2: RTDB telemetri + sinkronisasi jadwal otomatis',
    'useEffect(() => {',
    '  if (!claimedDeviceId || !rtdb) return;',
    '',
    '  // 2a. Listener telemetri live (isOnline, weight, stock, servo, time)',
    '  const unsubTelemetry = onValue(',
    '    ref(rtdb, `devices/${claimedDeviceId}`),',
    '    (snap) => setRtdbDeviceData(snap.exists() ? snap.val() : null)',
    '  );',
    '',
    '  // 2b. Listener scheduledFeedLog -- sinkronisasi jadwal ke Firestore',
    '  const unsubSchedLog = onValue(',
    '    ref(rtdb, `devices/${claimedDeviceId}/scheduledFeedLog`),',
    '    async (snap) => {',
    '      if (!snap.exists()) return;',
    '      const data = snap.val();',
    '      if (data.processed !== false || !catRef.current?.id) return;',
    '      // Set processed=true DULUAN -- cegah double-log',
    '      await set(ref(rtdb, `.../scheduledFeedLog/processed`), true);',
    '      // Buat Firestore feedingLog dengan notes="scheduled"',
    '      await addDoc(collection(db, "feedingLogs"), {',
    '        catId: catRef.current.id, deviceId: claimedDeviceId,',
    '        timestamp: data.ts ?? Date.now(),',
    '        amountRequested: data.amount, amountDispensed: data.amount,',
    '        status: "success", notes: "scheduled",',
    '      });',
    '    }',
    '  );',
    '  return () => { unsubTelemetry(); unsubSchedLog(); };',
    '}, [claimedDeviceId]);',
    '',
    '// Deteksi isRtdbLive: data ESP32 harus < 2 menit yang lalu',
    'const isRtdbLive = (() => {',
    '  void staleCheck;  // reactive: re-evaluasi tiap 30 detik',
    '  const timeStr = rtdbDeviceData?.time;  // "HH:MM:SS" dari NTP ESP32',
    '  if (!timeStr) return false;',
    '  const [h,m,s] = timeStr.split(":").map(Number);',
    '  const now = new Date();',
    '  const devSec = h*3600 + m*60 + s;',
    '  const nowSec = now.getHours()*3600 + now.getMinutes()*60 + now.getSeconds();',
    '  let diff = Math.abs(nowSec - devSec);',
    '  if (diff > 43200) diff = 86400 - diff;  // handle tengah malam',
    '  return diff < 120;  // toleransi 2 menit',
    '})();',
    '',
    '// Merge Firestore + RTDB menjadi satu DeviceStatus',
    'const device: DeviceStatus | null = fsDevice ? {',
    '  ...fsDevice,  // claim metadata dari Firestore',
    '  isOnline:             rtdbDeviceData?.isOnline === true && isRtdbLive,',
    '  foodStockLevel:       rtdbDeviceData?.foodStock ?? fsDevice.foodStockLevel,',
    '  currentWeightOnScale: Math.round(rtdbDeviceData?.weight ?? 0),',
    '  servoStatus:          rtdbDeviceData?.servoStatus ?? fsDevice.servoStatus,',
    '  lastPulse:            parseRtdbTime(rtdbDeviceData?.time) || fsDevice.lastPulse,',
    '} : null;',
], caption='Kode 4.5 useCatData.ts — Data Fusion Firestore + RTDB')

add_heading('4.4.5 FeedingControl.tsx — Kirim Perintah ke ESP32', size=11, bold=True)
add_body(
    'FeedingControl adalah satu-satunya komponen yang berwenang menulis perintah '
    'ke Firebase RTDB node /command/. Komponen ini juga mengelola jadwal makan '
    'di Firestore dan mengimplementasikan fitur Smart Adjustment '
    '(redistribusi sisa porsi hari ini jika kucing sudah makan manual).'
)
add_code([
    '// src/components/control/FeedingControl.tsx -- fungsi kirim perintah',
    '',
    '// Kirim perintah feed manual ke ESP32 via Firebase RTDB',
    'async function handleManualFeed(grams: number) {',
    '  if (!device?.id || !rtdb) return;',
    '',
    '  // 1. Tulis perintah ke RTDB -- ESP32 polling tiap 1 detik',
    '  await set(ref(rtdb, `devices/${device.id}/command`), {',
    '    type:   "feed",',
    '    amount: grams,',
    '    status: "pending",   // ESP32 ubah jadi "executing" lalu "done"',
    '  });',
    '',
    '  // 2. Buat feedingLog di Firestore dengan notes="manual"',
    '  await addDoc(collection(db, "feedingLogs"), {',
    '    catId:           cat.id,',
    '    deviceId:        device.id,',
    '    timestamp:       Date.now(),',
    '    amountRequested: grams,',
    '    amountDispensed: grams,',
    '    status:          "success",',
    '    notes:           "manual",',
    '  });',
    '}',
    '',
    '// Simpan jadwal makan ke Firestore /cats/{catId}/feedingSchedule',
    '// dan sinkronisasi ke RTDB /devices/{id}/schedule untuk dibaca ESP32',
    'async function saveSchedule(slots: FeedingScheduleSlot[]) {',
    '  await updateDoc(doc(db, "cats", cat.id), { feedingSchedule: slots });',
    '  await set(ref(rtdb, `devices/${device.id}/schedule`), {',
    '    enabled: true,',
    '    dailyLimitReached: false,',
    '    slots: slots.map(s => ({ time: s.time, amount: s.amount, active: s.active })),',
    '  });',
    '}',
], caption='Kode 4.6 FeedingControl.tsx — Kirim Perintah via RTDB & Simpan Jadwal')

add_heading('4.4.6 types.ts — Kontrak Data Seluruh Sistem', size=11, bold=True)
add_body(
    'File types.ts mendefinisikan seluruh interface dan enum TypeScript yang '
    'menjadi kontrak data antar komponen, Firebase, dan ESP32. '
    'Dengan TypeScript, compiler langsung mendeteksi jika ada field yang salah '
    'sebelum kode dijalankan — ini krusial untuk sistem IoT di mana data '
    'datang dari sumber eksternal (ESP32 via RTDB) yang tidak bisa dikontrol runtime-nya.'
)
add_code([
    '// src/types.ts -- interface utama',
    '',
    'export enum UserRole {',
    '  SUPER_ADMIN = "SUPER_ADMIN",',
    '  USER = "USER",',
    '}',
    '',
    'export interface CatProfile {',
    '  id: string;',
    '  ownerId: string;',
    '  name: string;',
    '  photoUrl?: string;',
    '  gender: "male" | "female";',
    '  age: number;',
    '  weight: number;',
    '  isSterilized: boolean;',
    '  bodyCondition: 1 | 2 | 3 | 4 | 5;  // BCS skala 1-5',
    '  kiloCaloriesPerKg: number;',
    '  dailyCalorieTarget: number;',
    '  dailyGramTarget: number;',
    '  feedingSchedule: FeedingScheduleSlot[];',
    '  profileUpdatedAt?: number;          // unix ms -- trigger reset todayTotal',
    '  dailyLimitReachedDate?: string;     // "YYYY-MM-DD" -- blokir jadwal otomatis',
    '  dailyAdjustments?: DailyAdjustment; // Smart Adjustment sisa hari ini',
    '}',
    '',
    'export interface DeviceStatus {',
    '  id: string;',
    '  isOnline: boolean;             // Merge dari RTDB isOnline && isRtdbLive',
    '  lastPulse: number;             // unix ms dari field time RTDB',
    '  foodStockLevel: number;        // % 0-100 dari HC-SR04',
    '  currentWeightOnScale: number;  // gram dari HX711',
    '  servoStatus: "idle" | "active" | "jammed";',
    '  calibrationFactor: number;     // loadCellFactor untuk HX711',
    '}',
    '',
    'export interface FeedingLog {',
    '  id: string;',
    '  catId: string;',
    '  deviceId?: string;',
    '  timestamp: number;             // unix ms',
    '  amountRequested: number;       // gram yang diminta',
    '  amountDispensed: number;       // gram yang benar-benar tertuang',
    '  status: "success" | "failed" | "warning";',
    '  notes?: string;                // "manual" | "auto" | "scheduled"',
    '}',
], caption='Kode 4.7 types.ts — Interface Utama Sistem PawfectCare')

add_heading('4.4.7 Ringkasan Ketergantungan Antar File', size=11, bold=True)
add_diagram([
    '  DEPENDENCY GRAPH APLIKASI WEB PAWFECTCARE',
    '',
    '  firebase.ts',
    '     |-- auth, db, rtdb, storage (diekspor ke semua file)',
    '     |',
    '  AuthContext.tsx',
    '     |-- onAuthStateChanged (firebase/auth)',
    '     |-- onSnapshot /users/{uid} (Firestore)',
    '     |-- mengekspos: useAuth() -> { user, profile, isAdmin, loading }',
    '     |',
    '  useCatData.ts',
    '     |-- useAuth() -> targetOwnerId',
    '     |-- onSnapshot /cats, /devices, /catProfileHistory (Firestore)',
    '     |-- onValue /devices/{id} + /scheduledFeedLog (RTDB)',
    '     |-- onSnapshot /feedingLogs (Firestore)',
    '     |-- mengekspos: { cat, device, feedingLogs, profileHistory }',
    '     |',
    '  App.tsx',
    '     |-- useAuth() -> routing berbasis role & state login',
    '     |-- merender komponen sesuai currentTab',
    '     |',
    '  Dashboard, Analytics, History, Notifications',
    '     |-- useCatData() -> data siap pakai',
    '     |',
    '  FeedingControl',
    '     |-- useCatData() -> cat, device',
    '     |-- set() ke RTDB /command -> perintah ke ESP32',
    '     |-- updateDoc() ke Firestore /cats -> jadwal',
    '     |-- addDoc() ke Firestore /feedingLogs -> log manual',
], caption='Gambar 4.1 Dependency Graph Aplikasi Web PawfectCare')

add_heading('4.5 Implementasi Fitur Aplikasi Web', size=12, bold=True)
add_table_styled(
    ['Halaman / Fitur', 'Role', 'Deskripsi Singkat'],
    [
        ['Dashboard', 'Admin + User', 'Progress harian, stok pakan, berat mangkuk, status servo, grafik konsumsi 4 jam dan mingguan'],
        ['Feeding Control', 'Admin', 'Input gram manual, trigger pakan instan, manajemen jadwal (tambah/edit/hapus), smart adjustment'],
        ['Cat Profile', 'Admin', 'Input/edit profil kucing, upload foto, kalkulasi kalori otomatis, riwayat profil'],
        ['Analytics', 'Admin + User', 'Rata-rata harian, grafik 7 hari, distribusi waktu makan (pie chart), % terjadwal vs manual, kalender bulanan'],
        ['Feeding History', 'Admin + User', 'Tabel riwayat log lengkap dengan filter tanggal'],
        ['Notifications', 'Admin + User', 'Alert stok rendah, overfeeding, perangkat offline'],
        ['Education', 'Semua', 'Artikel nutrisi kucing, BCS guide, tips perawatan'],
        ['Device Settings', 'Admin', 'Kalibrasi load cell (update loadCellFactor ke RTDB), info device, lepas device'],
        ['User Settings', 'Admin', 'Manajemen akun USER yang dapat monitoring'],
    ],
    caption='Tabel 4.3 Fitur Aplikasi Web PawfectCare'
)

add_heading('4.5 Fitur Edukasi', size=12, bold=True)
add_body(
    'Halaman Edukasi menyediakan konten ilmiah berbasis jurnal veteriner internasional '
    'yang dapat diakses oleh semua pengguna (admin maupun user monitoring). '
    'Halaman ini berisi 4 artikel utama, formula perhitungan pakan interaktif, '
    'dan daftar 14 referensi jurnal peer-reviewed yang menjadi landasan sistem.'
)

add_heading('4.5.1 Artikel Edukasi', size=11, bold=True)
add_table_styled(
    ['Artikel', 'Kategori', 'Temuan Utama', 'Referensi Kunci'],
    [
        ['FLUTD: Bahaya di Balik Pakan Berlebih',
         'Kesehatan',
         'Diet terapeutik mengurangi rekurensi FIC hingga 7,89x; pakan basah 11% vs kering 39% rekurensi',
         'Naarden & Corbee (2019); Bartges & Kirk (2007); Systematic Review NZ Vet J (2025)'],
        ['Formula Ilmiah di Balik Takaran Pakan',
         'Nutrisi & Sains',
         'Formula 4-faktor: Gram = (Fm x Fg x Fo x Fa x RER) / E; default 6x/hari berbasis pola makan alami kucing',
         'NRC (2006); WSAVA (2021); Bradshaw (2006); Zoran (2002)'],
        ['Seberapa Akurat Timbangan & Sensor?',
         'Teknologi',
         'HX711 24-bit akurasi 96,59-98,97% setelah kalibrasi; filter EMA alpha=0,15 eliminasi noise +/-1g',
         'Sri Hartanto (2022); Tim UMY (2022)'],
        ['Hidrasi: Pertahanan Utama Kesehatan Kucing',
         'Pencegahan',
         'Target 40-60 ml/kg/hari; pakan kering hanya 10% air; cat fountain meningkatkan konsumsi 2x',
         'WSAVA (2021); Bartges & Kirk (2007)'],
    ],
    caption='Tabel 4.4 Artikel Edukasi PawfectCare'
)

add_heading('4.5.2 Formula Perhitungan Pakan Lengkap (4 Faktor)', size=11, bold=True)
add_body(
    'Halaman edukasi menampilkan breakdown interaktif formula perhitungan pakan '
    'berbasis NRC 2006 dan WSAVA 2021 dengan 4 faktor penyesuaian:'
)
p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_formula = p_formula.add_run('Gram/hari = (Fm x Fg x Fo x Fa x 70 x BB^0.75) / E')
r_formula.font.name = 'Courier New'; r_formula.font.size = Pt(13); r_formula.font.bold = True

add_table_styled(
    ['Faktor', 'Nama', 'Nilai', 'Sumber'],
    [
        ['Fm', 'Maintenance (Umur + Steril)',
         'Kitten 0-6bln=3,0 | Kitten 6-12bln=2,0 | Dewasa steril=1,2 | Dewasa tdk steril=1,6 | Senior steril=1,1 | Geriatri>12th=0,8',
         'WSAVA 2021, NRC 2006'],
        ['Fg', 'Jenis Kelamin (opsional)',
         'Jantan=1,0 (referensi) | Betina=0,9 (koreksi risiko +BB)',
         'Wolfsheimer 1994 (Kirk\'s CVT XII)'],
        ['Fo', 'Body Condition Score',
         'BCS 1-2 (kurus)=1,2 | BCS 3 (ideal)=1,0 | BCS 4 (gemuk)=0,9 | BCS 5 (obesitas)=0,8',
         'WSAVA 2021, NRC 2006'],
        ['Fa', 'Aktivitas Fisik',
         'Sangat tdk aktif=0,8 | Tdk aktif=0,9 | Normal=1,0 | Aktif=1,1 | Sangat aktif=1,3',
         'Hand et al. (2010)'],
        ['E', 'Metabolizable Energy pakan',
         'Default 4,0 kcal/gram (rata-rata pakan kering 3700-4200 kcal/kg, NRC 2006)',
         'NRC (2006), Case et al. (2011)'],
        ['BB', 'Berat Badan aktual',
         'Dalam kilogram; untuk BCS 5 (obesitas) gunakan berat ideal',
         'WSAVA 2021'],
    ],
    caption='Tabel 4.5 Penjelasan 4 Faktor Formula Perhitungan Pakan'
)

add_body(
    'Alasan default 6x pemberian per hari: kucing secara alami makan 8-16x/hari '
    'dalam porsi kecil (Bradshaw, 2006). Kapasitas lambung kucing hanya 40-60 mL '
    '(Zoran, 2002), sehingga 6x per hari adalah titik tengah praktis yang mendistribusikan '
    'kalori merata, mencegah overfeeding, dan menjaga pH urine lebih stabil (Buffington et al., 2006).'
)

add_heading('4.5.3 Referensi Jurnal di Halaman Edukasi', size=11, bold=True)
add_table_styled(
    ['No', 'Penulis (Tahun)', 'Judul Singkat', 'Relevansi untuk Sistem'],
    [
        ['1', 'Naarden & Corbee (2019)', 'Effect of therapeutic urinary stress diet on FIC recurrence', 'Diet terkontrol kurangi risiko FIC 7,89x -- justifikasi kontrol porsi'],
        ['2', 'Bartges & Kirk (2007)', 'Nutrition and Urinary Tract Disease in Cats', 'Pakan basah 11% vs kering 39% rekurensi -- edukasi jenis pakan'],
        ['3', 'Buffington et al. (2006)', 'Multimodal environmental modification for idiopathic cystitis', 'Multiple small meals kurangi stres & mendukung manajemen FIC'],
        ['4', 'Systematic Review NZVetJ (2025)', 'Evidence base for FIC management strategies', 'Bukti terbaru 2025 -- strategi manajemen FLUTD paling valid'],
        ['5', 'Witzel-Rollins et al. (2022)', 'Unsupervised weight loss: automatic feeder & meal frequency', 'Automatic feeder: 83,2% kucing capai BCS ideal vs 0% free-feeding'],
        ['6', 'NRC (2006)', 'Nutrient Requirements of Dogs and Cats', 'Tabel life-stage multiplier RER -- dasar nilai Fm & Fg'],
        ['7', 'WSAVA (2021)', 'WSAVA Global Nutrition Guidelines', 'Standar internasional nutrisi -- nilai DER multiplier & BCS'],
        ['8', 'Hand et al. (2010)', 'Small Animal Clinical Nutrition, 5th Ed.', 'Sumber nilai Fa (faktor aktivitas) & panduan DER klinis'],
        ['9', 'Case et al. (2011)', 'Canine and Feline Nutrition, 3rd Ed.', 'ME pakan kering 3500-4500 kcal/kg -- dasar default E=4 kcal/g'],
        ['10', 'Olson (2022)', 'Energy Calculations: Caloric Intake for Patients', 'Konfirmasi klinis formula RER=70xBW^0,75 pada praktek modern'],
        ['11', 'Bradshaw (2006)', 'Evolutionary basis for feeding behavior of cats', 'Kucing makan 8-16x/hari -- dasar default 6x feeding'],
        ['12', 'Zoran (2002)', 'The carnivore connection to nutrition in cats', 'Lambung kucing 40-60 mL -- anatomi multiple small meals'],
        ['13', 'Sri Hartanto (2022)', 'Timbangan Digital Load Cell HX711 20kg', 'Validasi implementasi HX711 -- akurasi & kalibrasi'],
        ['14', 'Tim UMY (2022)', 'Analisis Akurasi HX711 Timbangan Bayi', 'Akurasi 96,59-98,97% setelah kalibrasi -- dasar klaim +/-1,5g'],
    ],
    caption='Tabel 4.6 Daftar 14 Referensi Jurnal di Halaman Edukasi'
)

add_heading('4.6 Pengujian dan Evaluasi', size=12, bold=True)

add_heading('4.6.1 Pengujian Akurasi Sensor Berat (HX711)', size=11, bold=True)
add_table_styled(
    ['Percobaan', 'Target (g)', 'Referensi (g)', 'HX711 (g)', 'Selisih (g)', 'Status'],
    [
        ['1', '30', '30.2', '30', '0.2', 'PASS'],
        ['2', '50', '49.8', '50', '0.2', 'PASS'],
        ['3', '75', '74.5', '75', '0.5', 'PASS'],
        ['4', '100', '100.3', '100', '0.3', 'PASS'],
        ['5', '40', '39.7', '40', '0.3', 'PASS'],
    ],
    caption='Tabel 4.7 Pengujian Akurasi Sensor Berat HX711'
)
add_body('Rata-rata selisih: 0.3 gram. Akurasi: 99.4%. Memenuhi toleransi +/- 1.5 gram.')

add_heading('4.6.2 Pengujian Akurasi Dispensasi Pakan', size=11, bold=True)
add_table_styled(
    ['Percobaan', 'Target (g)', 'Tertuang (g)', 'Selisih (g)', 'Fase Fine Tuning', 'Status'],
    [
        ['1', '50', '49.5', '0.5', '0 pulse', 'PASS'],
        ['2', '50', '51.2', '1.2', '2 pulse', 'PASS'],
        ['3', '30', '28.8', '1.2', '3 pulse', 'PASS'],
        ['4', '75', '74.3', '0.7', '1 pulse', 'PASS'],
        ['5', '100', '98.9', '1.1', '2 pulse', 'PASS'],
        ['6', '30', '31.4', '1.4', '0 pulse', 'PASS'],
    ],
    caption='Tabel 4.8 Pengujian Akurasi Dispensasi Pakan'
)
add_body('Rata-rata selisih: 1.0 gram. Semua dalam toleransi +/- 1.5 gram.')

add_heading('4.6.3 Pengujian Latensi Data Real-time', size=11, bold=True)
add_table_styled(
    ['Skenario', 'Latensi Rata-rata', 'Latensi Maks', 'Keterangan'],
    [
        ['ESP32 -> Firebase RTDB', '< 500 ms', '1200 ms', 'Kirim sensor tiap 2 detik'],
        ['Perintah web -> ESP32 eksekusi', '< 2000 ms', '3000 ms', 'Polling checkCommand() tiap 1 detik'],
        ['Firebase -> Dashboard (onSnapshot)', '< 200 ms', '600 ms', 'WebSocket listener real-time'],
        ['Jadwal terpicu -> servo buka', '0 - 1000 ms', '1000 ms', 'checkScheduledFeeding tiap 1 detik'],
    ],
    caption='Tabel 4.9 Pengujian Latensi Sistem'
)

add_heading('4.6.4 Pengujian Sensor Stok Pakan (HC-SR04)', size=11, bold=True)
add_table_styled(
    ['Kondisi Wadah', 'Jarak (cm)', 'Stok (%)', 'Nilai Sebenarnya'],
    [
        ['Penuh', '3.1', '98%', '~100%'],
        ['3/4 penuh', '6.5', '74%', '~75%'],
        ['1/2 penuh', '10.1', '49%', '~50%'],
        ['1/4 penuh', '13.8', '23%', '~25%'],
        ['Hampir kosong', '16.4', '5%', '~0-10%'],
    ],
    caption='Tabel 4.10 Pengujian Sensor Stok Pakan Feeder 1'
)

add_heading('4.6.5 Pengujian Jadwal Otomatis', size=11, bold=True)
add_table_styled(
    ['Slot', 'Waktu Setting', 'Waktu Aktual', 'Selisih', 'Status'],
    [
        ['Pagi', '07:00', '07:00', '0 detik', 'PASS'],
        ['Siang', '12:00', '12:00 - 12:01', '0 - 60 detik', 'PASS'],
        ['Sore', '17:00', '17:00', '0 detik', 'PASS'],
        ['Malam', '21:00', '21:00 - 21:01', '0 - 60 detik', 'PASS'],
    ],
    caption='Tabel 4.11 Pengujian Jadwal Otomatis'
)
add_body(
    'Toleransi maksimal 60 detik (bukan 1 menit) karena checkScheduledFeeding() '
    'berjalan setiap 1 detik — keterlambatan hanya bisa terjadi jika ESP32 '
    'sedang menjalankan proses feeding saat waktu jadwal tiba.'
)
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  BAB V — PENUTUP
# ═══════════════════════════════════════════════════════════════════════════
add_heading('BAB V', size=14, center=True, space_before=0)
add_heading('PENUTUP', size=14, center=True, space_before=0)

add_heading('5.1 Kesimpulan', size=12, bold=True)
add_number(
    'Sistem PawfectCare berhasil diimplementasikan dengan dua perangkat ESP32: '
    'Feeder 1 sebagai unit dispensasi utama (HC-SR04 + HX711 + Servo + OLED SH1106) '
    'dan Feeder 2 sebagai unit monitoring tambahan (HC-SR04 + OLED SSD1306).'
)
add_number(
    'Algoritma dispensasi dua fase (Bulk + Fine Tuning) dengan parameter aktual '
    'DISPENSE_BUFFER=7g, PULSE_MS=80ms, SETTLE_MS=2500ms mencapai akurasi rata-rata '
    '1.0 gram dari target, dengan semua hasil dalam toleransi +/-1.5 gram.'
)
add_number(
    'Komunikasi real-time antara ESP32 dan aplikasi web melalui Firebase RTDB '
    '(latensi < 500ms) dan Firestore (onSnapshot) berhasil dengan latensi perintah '
    'end-to-end rata-rata di bawah 2 detik.'
)
add_number(
    'Mekanisme deteksi status online menggunakan dua lapisan: isOnline flag di RTDB '
    'dan pengecekan freshness data (isRtdbLive, toleransi 2 menit) di aplikasi web, '
    'sehingga status device akurat meski koneksi tidak stabil.'
)
add_number(
    'Sinkronisasi log jadwal otomatis dari RTDB (scheduledFeedLog) ke Firestore '
    'feedingLogs berhasil tanpa duplikasi menggunakan mekanisme processed flag.'
)
add_number(
    'Sistem kalkulasi RER berhasil menghitung target kalori dan gram pakan secara '
    'otomatis sesuai profil biologis kucing (berat, usia, sterilisasi, BCS, aktivitas).'
)

add_heading('5.2 Saran Pengembangan', size=12, bold=True)

saran = [
    ('1. Aplikasi Mobile Native',
     'Kembangkan ke React Native/Flutter untuk push notification yang lebih andal, '
     'terutama di iOS yang membatasi notifikasi latar belakang dari web app.'),
    ('2. Computer Vision (Deteksi Kucing Makan)',
     'Tambahkan ESP32-CAM dengan model YOLOv8-nano untuk mendeteksi secara visual '
     'apakah kucing benar-benar mendatangi mangkuk. Sistem saat ini hanya mendeteksi '
     'penurunan berat mangkuk yang bisa disebabkan faktor lain.'),
    ('3. Multi-Kucing per Device',
     'Tambahkan identifikasi kucing via microchip RFID di kalung agar porsi pakan '
     'dapat dibedakan per individu dalam satu perangkat.'),
    ('4. Offline Mode',
     'Implementasikan local storage di ESP32 (SPIFFS/LittleFS) agar jadwal tetap '
     'berjalan meskipun koneksi internet terputus.'),
    ('5. Sensor Kualitas Pakan',
     'Tambahkan sensor kelembapan (DHT22) di dalam hopper untuk mendeteksi '
     'apakah pakan mulai lembap atau berjamur.'),
    ('6. Timbangan Berat Badan Kucing',
     'Tambahkan platform timbangan terpisah untuk mengukur berat badan kucing secara '
     'rutin dan menyesuaikan target pakan secara otomatis.'),
    ('7. Ekspor Laporan Veteriner',
     'Kembangkan fitur ekspor laporan nutrisi bulanan PDF untuk keperluan konsultasi '
     'dokter hewan.'),
]

for title, desc in saran:
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(8)
    p_t.paragraph_format.space_after  = Pt(2)
    p_t.paragraph_format.left_indent  = Cm(0.5)
    set_run_font(p_t.add_run(title), size=12, bold=True)
    b = add_body(desc, indent=False)
    b.paragraph_format.left_indent = Cm(0.5)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  DAFTAR PUSTAKA
# ═══════════════════════════════════════════════════════════════════════════
add_heading('DAFTAR PUSTAKA', size=14, center=True, space_before=0)
add_body(
    'Daftar pustaka disusun secara alfabetis berdasarkan nama penulis pertama. '
    'Referensi mencakup sumber IoT/teknologi dan referensi veteriner/nutrisi '
    'yang menjadi landasan ilmiah fitur edukasi sistem PawfectCare.',
    indent=False
)
doc.add_paragraph()
refs = [
    # ── A ──
    'Atzori, L., Iera, A., & Morabito, G. (2010). The Internet of Things: A survey. '
    'Computer Networks, 54(15), 2787-2805.',

    # ── B ──
    'Bartges, J. & Kirk, C. (2007). Nutrition and Urinary Tract Disease in Cats: Myths and Legends. '
    'Journal of Feline Medicine and Surgery, 9(6), 487-490. https://doi.org/10.1016/S1098-612X-07-00199-4',

    'Bradshaw, J.W.S. (2006). The evolutionary basis for the feeding behavior of domestic dogs and cats. '
    'Journal of Nutrition, 136(7 Suppl), 1927S-1931S. https://doi.org/10.1093/jn/136.7.1927S',

    'Buffington, C.A.T., Westropp, J.L., Chew, D.J., & Bolus, R.R. (2006). Clinical evaluation of '
    'multimodal environmental modification in the management of cats with idiopathic cystitis. '
    'Journal of Feline Medicine and Surgery, 8(4), 261-268. https://doi.org/10.1016/j.jfms.2006.02.002',

    # ── C ──
    'Case, L.P., Daristotle, L., Hayek, M.G., & Raasch, M.F. (2011). Canine and Feline Nutrition: '
    'A Resource for Companion Animal Professionals, 3rd Edition. Mosby/Elsevier.',

    # ── E ──
    'Espressif Systems. (2023). ESP32 Technical Reference Manual v5.1. Espressif Systems.',

    # ── F ──
    'Firebase. (2024). Firebase Documentation - Realtime Database & Cloud Firestore. '
    'Google LLC. https://firebase.google.com/docs',

    # ── H ──
    'Hand, M.S., Thatcher, C.D., Remillard, R.L., Roudebush, P., & Novotny, B.J. (2010). '
    'Small Animal Clinical Nutrition, 5th Edition. Mark Morris Institute.',

    'Hartanto, S. (2022). Rancang Bangun Timbangan Digital Load Cell Berkapasitas 20 kg '
    'Berbasis Modul HX711. Jurnal Elektro, Universitas Krisnadwipayana, hal. 21-26. '
    'https://jurnalteknik.unkris.ac.id/index.php/jie/article/view/561',

    # ── L ──
    'Latha, V., & Kumar, A. (2021). Real-time Cloud Integration using Firebase for IoT Applications. '
    'International Journal of Advanced Computer Science and Applications, 12(4), 201-208.',

    # ── M ──
    'Maulana, R., Santoso, A., & Wibowo, T. (2021). Rancang Bangun Smart Pet Feeder Berbasis '
    'Android dan NodeMCU. Jurnal Teknologi Informasi dan Komunikasi, 10(2), 45-54.',

    # ── N ──
    'Naarden, B. & Corbee, R.J. (2019). The effect of a therapeutic urinary stress diet on the '
    'short-term recurrence of feline idiopathic cystitis. Veterinary Medicine and Science, '
    '6(1), 32-38. https://doi.org/10.1002/vms3.197',

    'National Research Council (NRC). (2006). Nutrient Requirements of Dogs and Cats. '
    'National Academies Press, Washington DC. '
    'https://nap.nationalacademies.org/catalog/10668/nutrient-requirements-of-dogs-and-cats',

    'Ngo, T., & Kamata, S. (2020). HX711 Load Cell Amplifier Calibration for High-Accuracy '
    'Weight Measurement in Embedded Systems. Sensors and Actuators A: Physical, 315, 112303.',

    # ── O ──
    'Olson, M. (LVT, VTS). (2022). Energy Calculations: Gauging the Proper Caloric Intake for Patients. '
    "Today's Veterinary Nurse, Summer 2022. "
    'https://todaysveterinarynurse.com/nutrition/veterinary-energy-calculations-and-proper-caloric-intake/',

    # ── P ──
    'Prasetyo, D., Kurniawan, H., & Nugroho, B. (2022). Sistem Monitoring Pakan Kucing Real-time '
    'Menggunakan Raspberry Pi dan Firebase. Jurnal Ilmu Komputer dan Informatika, 8(1), 23-32.',

    # ── R ──
    'React Documentation. (2024). React - The library for web and native user interfaces. '
    'Meta. https://react.dev',

    # ── T ──
    'Taylor & Francis (Systematic Review). (2025). Understanding the current evidence base for the '
    'commonly recommended management strategies for recurrent feline idiopathic cystitis: a '
    'systematic review. New Zealand Veterinary Journal. '
    'https://doi.org/10.1080/00480169.2025.2477542',

    'Tim Peneliti UMY. (2022). Analisis Akurasi Modul Amplifier HX711 untuk Timbangan Bayi. '
    'Medika Teknika: Jurnal Teknik Elektromedik Indonesia, Universitas Muhammadiyah Yogyakarta. '
    'https://journal.umy.ac.id/index.php/mt/article/view/15148',

    # ── W ──
    'Witzel-Rollins, A., Murphy, M., Springer, C.M., Moyers, T.D., & Albright, J.D. (2022). '
    'Unsupervised weight loss in multi-cat households: effects of an automatic feeder and meal '
    'frequency. Journal of Feline Medicine and Surgery. '
    'https://doi.org/10.1177/1098612X221105046',

    'WSAVA Global Nutrition Committee. (2021). WSAVA Global Nutrition Guidelines & Toolkit. '
    'World Small Animal Veterinary Association. '
    'https://wsava.org/global-guidelines/global-nutrition-guidelines/',

    # ── Z ──
    'Zhao, X., Li, T., & Wang, Y. (2019). Design and Implementation of an Automated Pet Feeder '
    'with IoT Monitoring. Proceedings of IEEE IoT Conference 2019, 112-117.',

    'Zoran, D.L. (2002). The carnivore connection to nutrition in cats. Journal of the American '
    'Veterinary Medical Association, 221(11), 1559-1567. '
    'https://doi.org/10.2460/javma.2002.221.1559',
]
for ref in refs:
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_after = Pt(6)
    p_ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_ref.paragraph_format.first_line_indent = Cm(-1.0)
    p_ref.paragraph_format.left_indent = Cm(1.0)
    set_run_font(p_ref.add_run(ref), size=11)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  LAMPIRAN
# ═══════════════════════════════════════════════════════════════════════════
add_heading('LAMPIRAN', size=14, center=True, space_before=0)

add_heading('Lampiran A — Konfigurasi Pin Lengkap', size=12, bold=True)
add_table_styled(
    ['Feeder', 'Pin ESP32', 'Fungsi', 'Komponen'],
    [
        ['Feeder 1', 'GPIO 12', 'OUTPUT — Trigger Ultrasonik', 'HC-SR04 TRIG'],
        ['Feeder 1', 'GPIO 14', 'INPUT — Echo Ultrasonik', 'HC-SR04 ECHO'],
        ['Feeder 1', 'GPIO 27', 'PWM — Kontrol Servo', 'Servo Motor'],
        ['Feeder 1', 'GPIO 32', 'INPUT — Data Load Cell', 'HX711 DT'],
        ['Feeder 1', 'GPIO 33', 'OUTPUT — Clock Load Cell', 'HX711 SCK'],
        ['Feeder 1', 'GPIO 21', 'I2C SDA', 'OLED SH1106'],
        ['Feeder 1', 'GPIO 22', 'I2C SCL', 'OLED SH1106'],
        ['Feeder 2', 'GPIO 17', 'OUTPUT — Trigger Ultrasonik', 'HC-SR04 TRIG'],
        ['Feeder 2', 'GPIO 5', 'INPUT — Echo Ultrasonik', 'HC-SR04 ECHO'],
        ['Feeder 2', 'GPIO 21', 'I2C SDA', 'OLED SSD1306'],
        ['Feeder 2', 'GPIO 22', 'I2C SCL', 'OLED SSD1306'],
    ]
)

add_heading('Lampiran B — Konstanta Firmware Feeder 1 (Nilai Aktual)', size=12, bold=True)
add_table_styled(
    ['Konstanta', 'Nilai Aktual', 'Penjelasan'],
    [
        ['MAX_SERVO_ANGLE', '20 derajat', 'Sudut buka servo maksimal'],
        ['PAN_WEIGHT_G', '3.0 g', 'Berat wadah pakan (untuk weightGross)'],
        ['DISPENSE_BUFFER_G', '7.0 g', 'Tutup servo X gram sebelum target'],
        ['DISPENSE_TOLERANCE_G', '1.5 g', 'Toleransi akhir dianggap berhasil'],
        ['DISPENSE_PULSE_MS', '80 ms', 'Durasi buka servo per-pulse fase fine'],
        ['DISPENSE_SETTLE_MS', '2500 ms', 'Waktu tunggu settling fase fine'],
        ['DISPENSE_MAX_PULSES', '6 kali', 'Maksimal pulse fase fine'],
        ['loadCellFactor', '404.0', 'Faktor kalibrasi HX711 (updateable via RTDB)'],
        ['Interval baca sensor', '250 ms', 'readFoodStock() + readWeight()'],
        ['Interval kirim RTDB', '2000 ms', 'sendDataToFirebase()'],
        ['Interval cek perintah', '1000 ms', 'checkCommand()'],
        ['Interval cek jadwal', '1000 ms', 'checkScheduledFeeding()'],
        ['Interval update OLED', '500 ms', 'drawMainScreen()'],
        ['Tare sample', '20 sample', 'scale.tare(20) saat inisialisasi'],
        ['fullDist Feeder 1', '3.0 cm', 'Jarak HC-SR04 saat hopper penuh'],
        ['emptyDist Feeder 1', '17.0 cm', 'Jarak HC-SR04 saat hopper kosong'],
        ['fullDist Feeder 2', '2.0 cm', 'Jarak HC-SR04 saat hopper penuh'],
        ['emptyDist Feeder 2', '13.0 cm', 'Jarak HC-SR04 saat hopper kosong'],
    ]
)

add_heading('Lampiran C — Tipe Data TypeScript (types.ts)', size=12, bold=True)
add_table_styled(
    ['Interface / Enum', 'Field Utama', 'Keterangan'],
    [
        ['UserRole (enum)', 'SUPER_ADMIN, USER', 'Dua peran pengguna'],
        ['UserProfile', 'uid, email, role, onboardingCompleted, claimedDeviceId', 'Profil akun'],
        ['CatProfile', 'name, gender, age, weight, isSterilized, bodyCondition, dailyGramTarget, feedingSchedule, dailyAdjustments', 'Profil kucing'],
        ['FeedingScheduleSlot', 'time (HH:mm), amount (gram), active', 'Satu slot jadwal'],
        ['FeedingLog', 'catId, deviceId, timestamp, amountRequested, amountDispensed, status, notes', 'Satu riwayat pemberian pakan'],
        ['DeviceStatus', 'isOnline, lastPulse, foodStockLevel, currentWeightOnScale, servoStatus, calibrationFactor', 'Status real-time perangkat'],
        ['CatProfileSnapshot', 'savedAt, endedAt + semua field CatProfile', 'Snapshot profil untuk history'],
        ['DailyAdjustment', 'date, manualTotal, slots:[{time, originalAmount, adjustedAmount}]', 'Smart adjustment sisa jadwal hari ini'],
    ]
)

add_heading('Lampiran D — Alur Lengkap Setup Firmware (Feeder 1)', size=12, bold=True)
add_diagram([
    'setup():',
    '  Serial.begin(115200)',
    '  pinMode(TRIG_PIN, OUTPUT); pinMode(ECHO_PIN, INPUT)',
    '  |',
    '  v',
    '  initOLED()   : Wire.begin(21,22) -> display.begin(0x3C) -> splash screen',
    '  |',
    '  v',
    '  connectWiFi(): WiFi.begin() -> tunggu max 40x500ms',
    '  |  gagal -> tampil error OLED -> return (tidak lanjut)',
    '  v',
    '  initNTP()    : timeClient.begin() -> forceUpdate()',
    '  |',
    '  v',
    '  initFirebase(): config API_KEY + DATABASE_URL + auth email/pass',
    '  |              -> Firebase.begin() -> tunggu max 20x500ms',
    '  |              -> jika ready: firebaseReady=true, registerOnDisconnect()',
    '  v',
    '  initServo()  : feederServo.attach(27) -> write(0)',
    '  |',
    '  v',
    '  initHX711()  : scale.begin(32,33) -> tunggu is_ready()',
    '  |              -> scale.set_scale(404.0)',
    '  |              -> delay(2000) -- stabilisasi',
    '  |              -> scale.tare(20) -- nol-kan 20 sample',
    '  v',
    '  readCalibrationFromRTDB(): ambil loadCellFactor dari RTDB',
    '  |                          -> scale.set_scale(nilai baru)',
    '  v',
    '  Tampil "SYSTEM READY" -> masuk loop()',
], caption='Lampiran D Alur Setup Firmware Feeder 1')

# ─────────────────────────────────────────────────────────────
#  SIMPAN
# ─────────────────────────────────────────────────────────────
output = r'd:\Tugas Akhir\Tugas-Akhir\Skripsi_PawfectCare_SmartCatFeeder.docx'
doc.save(output)
print('SUKSES! File:', output)
print('Paragraf:', len(doc.paragraphs))
print('Tabel   :', len(doc.tables))
